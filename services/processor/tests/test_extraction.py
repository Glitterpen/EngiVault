from pathlib import Path

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from app.extraction import extract_document
from app.validation import MIME_DOCX, MIME_DWG, MIME_PDF, MIME_XLSX


def test_pdf_units_use_one_based_pages(tmp_path: Path) -> None:
    source = tmp_path / "drawing.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    writer.add_blank_page(width=595, height=842)
    with source.open("wb") as stream:
        writer.write(stream)
    result = extract_document(source, MIME_PDF)
    assert [unit.page_number for unit in result.units] == [1, 2]
    assert result.preview_strategy == "original_pdf"


def test_docx_retains_paragraph_provenance(tmp_path: Path) -> None:
    source = tmp_path / "specification.docx"
    document = Document()
    document.add_paragraph("Design pressure is 120 barg.")
    document.add_paragraph("")
    document.add_paragraph("Material shall be carbon steel.")
    document.save(source)
    result = extract_document(source, MIME_DOCX)
    assert [unit.paragraph_number for unit in result.units] == [1, 3]
    assert result.preview_strategy == "rendered_pdf_pending"


def test_xlsx_retains_sheet_and_range(tmp_path: Path) -> None:
    source = tmp_path / "register.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "MDR"
    sheet.append(["Document", "Revision"])
    sheet.append(["EC-PIP-001", "C02"])
    workbook.save(source)
    result = extract_document(source, MIME_XLSX)
    assert result.units[1].sheet_name == "MDR"
    assert result.units[1].cell_range == "A2:B2"


def test_dwg_is_validation_only_until_cad_adapter() -> None:
    result = extract_document(Path("unused.dwg"), MIME_DWG)
    assert result.units == []
    assert result.preview_strategy == "cad_adapter_required"
