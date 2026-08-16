from __future__ import annotations

from pathlib import Path
from datetime import date
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RAW = ROOT / "screenshots"
PROCESSED = ROOT / "screenshots-processed"
OUTPUT = ROOT / "EngiCite_Marketing_Blueprint_Final.docx"
LOGO = ROOT.parents[1] / "apps" / "web" / "public" / "engicite-logo-transparent.png"

NAVY = "082345"
NAVY_2 = "10243E"
ORANGE = "ED7138"
ORANGE_LIGHT = "FCE9E0"
GREEN = "16785B"
GREEN_LIGHT = "E9F5F0"
INK = "14263D"
GRAY = "5E6B7A"
LIGHT = "F4F7FA"
MID = "DDE5EC"
WHITE = "FFFFFF"
RED = "A33B2E"
AMBER = "8A6200"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_table_borders(table, color=MID, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def paragraph_border_bottom(paragraph, color=ORANGE, size=18, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])
    set_run_font(run, size=8.5, color=GRAY)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, NAVY, 14, 7),
        ("Heading 3", 12, INK, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Kicker", "Figure Caption", "Small Note"):
        if name not in doc.styles:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    kicker = doc.styles["Kicker"]
    kicker.font.name = "Calibri"
    kicker.font.size = Pt(9)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(ORANGE)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(5)
    caption = doc.styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)
    note = doc.styles["Small Note"]
    note.font.name = "Calibri"
    note.font.size = Pt(8.5)
    note.font.color.rgb = rgb(GRAY)
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.15


def configure_sections(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("ENGICITE  /  MARKETING BLUEPRINT"), size=8.5, color=NAVY, bold=True)
    paragraph_border_bottom(p, color=MID, size=7, space=5)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6900, 2460], indent_dxa=0)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    p_left = left.paragraphs[0]
    p_left.paragraph_format.space_after = Pt(0)
    set_run_font(p_left.add_run("Confidential working draft  |  13 August 2026"), size=8, color=GRAY)
    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    set_run_font(p_right.add_run("Page "), size=8.5, color=GRAY)
    add_page_field(p_right)

    first_footer = section.first_page_footer
    pff = first_footer.paragraphs[0]
    pff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(pff.add_run("Internal marketing and sales enablement document"), size=8.5, color=GRAY)


def process_screenshots():
    PROCESSED.mkdir(parents=True, exist_ok=True)
    processed = []
    for source in sorted(RAW.glob("*.png")):
        image = Image.open(source).convert("RGB")
        is_landing = source.name.startswith("00-")
        if not is_landing:
            image = image.crop((0, 64, image.width, image.height))
        draw = ImageDraw.Draw(image)
        if is_landing:
            draw.rectangle((0, image.height - 55, 72, image.height), fill=(246, 248, 250))
        else:
            draw.rectangle((0, image.height - 55, 72, image.height), fill=(16, 36, 62))
        if source.name == "06-team-resources.png":
            draw.rectangle((286, 530, 790, 598), fill=(255, 255, 255))
            name_font = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 20)
            detail_font = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 14)
            draw.text((297, 542), "Process Engineer", fill=(20, 38, 61), font=name_font)
            draw.text((297, 570), "process.engineer@example.com - engineer", fill=(94, 107, 122), font=detail_font)
        target = PROCESSED / source.name
        image.save(target, format="PNG", optimize=True)
        processed.append(target)

    thumbs = []
    for path in processed:
        img = Image.open(path).convert("RGB")
        img.thumbnail((320, 230))
        thumbs.append((path.name, img.copy()))
    if thumbs:
        rows = (len(thumbs) + 3) // 4
        sheet = Image.new("RGB", (4 * 340, rows * 270), "white")
        sheet_draw = ImageDraw.Draw(sheet)
        for index, (name, img) in enumerate(thumbs):
            x = (index % 4) * 340 + 10
            y = (index // 4) * 270 + 10
            sheet.paste(img, (x, y + 22))
            sheet_draw.text((x, y), name, fill=(8, 35, 69))
        sheet.save(PROCESSED / "contact-sheet.png", optimize=True)


def add_kicker(doc, text):
    p = doc.add_paragraph(style="Kicker")
    p.add_run(text.upper())
    return p


def add_title(doc, text, size=23, color=NAVY, after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size=size, color=color, bold=True)
    return p


def add_deck(doc, text, after=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    set_run_font(p.add_run(text), size=11.5, color=GRAY)
    return p


def add_body(doc, text, bold_lead=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullet_rows(doc, items, fill=LIGHT, marker_color=ORANGE, widths=(480, 8880), font_size=10.3):
    table = doc.add_table(rows=len(items), cols=2)
    set_table_geometry(table, list(widths))
    remove_table_borders(table)
    for row, item in zip(table.rows, items):
        marker, body = row.cells
        shade_cell(marker, fill)
        shade_cell(body, fill)
        marker.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        body.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pm = marker.paragraphs[0]
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_after = Pt(0)
        set_run_font(pm.add_run("•"), size=13, color=marker_color, bold=True)
        pb = body.paragraphs[0]
        pb.paragraph_format.space_after = Pt(3)
        pb.paragraph_format.space_before = Pt(3)
        if isinstance(item, tuple):
            lead, rest = item
            set_run_font(pb.add_run(lead), size=font_size, color=INK, bold=True)
            set_run_font(pb.add_run(rest), size=font_size, color=INK)
        else:
            set_run_font(pb.add_run(item), size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metric_grid(doc, metrics, cols=2):
    rows = (len(metrics) + cols - 1) // cols
    widths = [9360 // cols] * cols
    table = doc.add_table(rows=rows, cols=cols)
    set_table_geometry(table, widths)
    remove_table_borders(table)
    index = 0
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=170, bottom=170, start=180, end=180)
            shade_cell(cell, LIGHT)
            if index < len(metrics):
                value, label, color = metrics[index]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                set_run_font(p.add_run(value), size=20, color=color, bold=True)
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_after = Pt(0)
                set_run_font(p2.add_run(label.upper()), size=8.5, color=GRAY, bold=True)
            index += 1
    return table


def add_matrix(doc, headers, rows, widths, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=8.6, color=WHITE, bold=True)
    for r_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            shade_cell(cell, WHITE if r_index % 2 == 0 else LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index == 0:
                set_run_font(p.add_run(str(value)), size=font_size, color=NAVY, bold=True)
            else:
                set_run_font(p.add_run(str(value)), size=font_size, color=INK)
        for i, cell in enumerate(cells):
            set_cell_width(cell, widths[i])
    return table


def add_status_chip(cell, text, fill, color):
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text.upper()), size=8.2, color=color, bold=True)


def page_break(doc):
    doc.add_page_break()


def add_screenshot(doc, filename, alt, width=6.5):
    path = PROCESSED / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))._inline
    inline.docPr.set("descr", alt)
    return p


def add_figure_caption(doc, number, text):
    p = doc.add_paragraph(style="Figure Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Figure {number}. {text} Live development interface captured 13 August 2026.")
    return p


def add_interface_page(doc, number, kicker, title, deck, screenshot, caption, proof_items):
    page_break(doc)
    add_kicker(doc, kicker)
    add_title(doc, title, size=21, after=6)
    add_deck(doc, deck, after=8)
    add_screenshot(doc, screenshot, f"EngiCite interface showing {caption.lower()}")
    add_figure_caption(doc, number, caption)
    add_bullet_rows(doc, proof_items, font_size=9.4)


def add_brand_callout(doc, heading, body, fill=NAVY, accent=ORANGE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=220, bottom=220, start=260, end=260)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run(heading.upper()), size=9, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.18
    set_run_font(p2.add_run(body), size=13, color=WHITE, bold=True)
    return table


def build_document():
    process_screenshots()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    doc.core_properties.title = "EngiCite Marketing Blueprint"
    doc.core_properties.subject = "Go-to-market positioning, capability map and live interface tour"
    doc.core_properties.author = "EngiCite Team"
    doc.core_properties.keywords = "EngiCite, engineering document control, MDR, oil and gas, SaaS"
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    # Cover
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(35)
    cover_logo = p_logo.add_run().add_picture(str(LOGO), width=Inches(4.9))._inline
    cover_logo.docPr.set("descr", "EngiCite company logo")
    add_kicker(doc, "Marketing blueprint / Product capability narrative")
    title = add_title(doc, "Control every deliverable.\nCite every decision.", size=31, after=13)
    for run in title.runs:
        run.font.color.rgb = rgb(NAVY)
    add_deck(doc, "A buyer-focused blueprint for positioning, demonstrating and launching EngiCite as the engineering document control and intelligence workspace for complex projects.", after=25)
    add_brand_callout(
        doc,
        "The product in one sentence",
        "EngiCite helps EPC, EPCIC and owner teams plan the MDR, coordinate discipline submissions, control every revision, issue client transmittals and retrieve project knowledge with verifiable evidence.",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("Prepared for"), size=9, color=ORANGE, bold=True)
    add_body(doc, "EngiCite leadership, commercial partners, pilot customers and sales representatives", after=3)
    p = doc.add_paragraph(style="Small Note")
    p.add_run("Version 1.0  |  13 August 2026  |  Internal working document")

    # Executive blueprint
    page_break(doc)
    add_kicker(doc, "Executive blueprint")
    add_title(doc, "A controlled operating system for engineering information", size=23)
    add_deck(doc, "EngiCite is not simply cloud storage. It connects the project plan, the MDR, discipline ownership, revision evidence, DCC acceptance, progress reporting and formal issue into one traceable workflow.")
    add_metric_grid(doc, [
        ("1", "controlled source of truth", NAVY),
        ("5", "role-specific workspaces", ORANGE),
        ("100%", "evidence-backed AI target", GREEN),
        ("7 years", "audit-retention design target", NAVY),
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    add_body(doc, "The commercial story", bold_lead="The commercial story")
    add_bullet_rows(doc, [
        ("Plan with confidence. ", "Turn an approved deliverables list or Excel MDR into a live, discipline-organised register with due dates."),
        ("Control submission quality. ", "Engineers upload only against assigned discipline deliverables; DCC previews and accepts before controlled issue."),
        ("See slippage early. ", "Progress, overdue work, discipline performance, weekly look-ahead and variance are visible before they become client problems."),
        ("Issue professionally. ", "Freeze accepted revisions into work packages and client transmittals with a controlled acknowledgement trail."),
        ("Find the proof. ", "Search and AI workflows return project evidence with document, revision and page-level provenance."),
    ])
    p = doc.add_paragraph(style="Small Note")
    p.paragraph_format.space_before = Pt(8)
    p.add_run("Blueprint reading guide: interface pages show the current development build; status pages distinguish demonstrated, configuration-dependent and release-readiness capabilities.")

    # Market problem and position
    page_break(doc)
    add_kicker(doc, "Market need and positioning")
    add_title(doc, "Engineering delivery is still fragmented across spreadsheets, folders and email", size=22)
    add_deck(doc, "The cost is not just time spent searching. It is weak ownership, late submissions, uncontrolled revisions, incomplete handover records and management decisions made without reliable evidence.")
    add_matrix(doc, ["Market pain", "What teams experience", "EngiCite response"], [
        ("MDR administration", "Manual entry, inconsistent naming and slow updates", "Excel import plus a structured, discipline-grouped register"),
        ("Discipline coordination", "Engineers submit through email or the wrong folder", "Assigned discipline workspaces and controlled upload routes"),
        ("Revision confidence", "Teams cannot easily confirm the latest accepted file", "Immutable revision history, preview and DCC acceptance"),
        ("Schedule visibility", "Missed dates become visible only during meetings", "Due dates, overdue alerts, progress and weekly look-ahead"),
        ("Client issue", "Transmittals are assembled manually with weak traceability", "Frozen issue sets, transmittal numbers and acknowledgement covers"),
        ("Knowledge retrieval", "Users search filenames rather than engineering evidence", "Full-text, semantic retrieval and citation-grounded answers"),
    ], [1800, 3300, 4260], font_size=8.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_brand_callout(doc, "Category position", "EngiCite is the engineering document control and intelligence platform for project teams that need disciplined delivery and defensible evidence - without forcing every user to become a document-control expert.")

    # Buyer map
    page_break(doc)
    add_kicker(doc, "Ideal customer profile and buyers")
    add_title(doc, "Start where document complexity and delivery risk are both high", size=22)
    add_deck(doc, "The strongest early market is project-based engineering work where multiple disciplines produce controlled deliverables against contractual dates.")
    add_matrix(doc, ["Priority segment", "Buying trigger", "Primary value"], [
        ("EPC / EPCIC contractors", "New FEED, DED, brownfield or construction project", "Faster mobilisation, fewer late deliverables, controlled client issue"),
        ("Owner-operators", "Multiple contractors and inconsistent handover records", "Oversight, auditability and a durable project knowledge base"),
        ("Engineering consultancies", "Growing project portfolio managed through Excel and shared drives", "Standardised workflows without heavyweight enterprise rollout"),
        ("Document-control service firms", "Need to manage several clients with clear separation", "Multi-tenant controls and repeatable DCC operations"),
    ], [2100, 3300, 3960], font_size=9.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    add_body(doc, "The buying committee", bold_lead="The buying committee")
    add_bullet_rows(doc, [
        ("Economic buyer - Project Director / Operations Director: ", "cares about risk, schedule, client confidence and predictable project reporting."),
        ("Champion - Lead Document Controller: ", "cares about MDR speed, revision accuracy, submission quality and formal transmission."),
        ("Technical approver - IT / Information Security: ", "cares about tenant isolation, access controls, auditability, retention and recovery."),
        ("Daily user - Discipline Engineer: ", "cares about a clear personal queue, due dates, simple upload and rapid feedback."),
    ], font_size=9.7)

    # Role operating model
    page_break(doc)
    add_kicker(doc, "Role-specific operating model")
    add_title(doc, "Each person sees the work they are accountable for", size=22)
    add_deck(doc, "EngiCite separates management oversight, document-control authority and engineering submission responsibility so that one login does not expose every control to every user.")
    add_matrix(doc, ["Role", "Primary workspace", "Core responsibility"], [
        ("Organisation administrator", "Portfolio and organisation oversight", "Create organisations and projects, maintain account structure, observe project health"),
        ("Project administrator / manager", "Project overview, health, resources and reports", "Set project brief and objectives, invite the DCC and project resources, manage schedule"),
        ("Document controller", "DCC control centre, MDR, reviews and transmissions", "Populate the MDR, assign disciplines, accept submissions and issue controlled records"),
        ("Discipline engineer", "My deliverables", "See assigned documents and dates, upload revisions only for authorised disciplines"),
        ("Viewer / stakeholder", "Authorised project documents", "Read, search, preview, compare and download according to policy"),
    ], [1900, 3000, 4460], font_size=8.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_brand_callout(doc, "Workflow in one line", "ADMIN SETS THE PROJECT  >  DCC BUILDS THE MDR  >  ENGINEERS SUBMIT  >  DCC ACCEPTS  >  MANAGEMENT TRACKS  >  CLIENT RECEIVES A CONTROLLED ISSUE")
    add_body(doc, "This operating model is a central differentiator: it makes the software feel simple to each user while preserving strong controls underneath.", after=0)

    # Interface tour
    add_interface_page(doc, 1, "Public product experience", "A precise promise before the buyer signs in", "The landing page positions EngiCite around controlled delivery and evidence, with a direct path to the secure workspace and a visual demonstration of citation-backed answers.", "00-landing-page.png", "Public landing page and product promise.", [
        ("Buyer clarity. ", "The headline communicates the two outcomes immediately: deliverable control and defensible decisions."),
        ("Product proof. ", "The interface mock demonstrates cited answers, register health and controlled documents without requiring a live login."),
    ])

    add_interface_page(doc, 2, "Multi-tenant workspace", "Secure organisation separation", "Users enter only the organisations and projects assigned to them. This creates a clean commercial story for contractors, owner teams and service companies managing multiple clients.", "02-organisation-workspace.png", "Authorised organisation workspace.", [
        ("Tenant awareness. ", "Organisation boundaries are explicit in the interface, not hidden inside a shared folder tree."),
        ("Simple navigation. ", "A user starts from authorised work instead of searching across the whole platform."),
    ])

    add_interface_page(doc, 3, "Document controller experience", "A DCC control centre built around daily priorities", "The home view gives the document controller a project-specific operating picture: total MDR deliverables, items awaiting review, unsubmitted work, overdue submissions and discipline resources.", "04-dcc-control-centre.png", "Document Control Centre with project-level operating metrics.", [
        ("Action before administration. ", "Invite an engineer, transmit documents or open the MDR from the first screen."),
        ("Exception-led work. ", "The DCC sees overdue and unsubmitted work instead of manually reconciling spreadsheets."),
    ])

    add_interface_page(doc, 4, "Master Document Register", "The MDR becomes a live project control surface", "Documents are organised by discipline and can be filtered, searched, registered individually or imported from Excel. The register communicates ownership, type, due date and status in one view.", "01-master-document-register.png", "Discipline-grouped Master Document Register.", [
        ("Fast mobilisation. ", "Import an existing MDR rather than recreating hundreds of deliverables manually."),
        ("Flexible organisation. ", "Switch between a flat register, discipline groups and discipline-specific filters."),
        ("Governed creation. ", "Document registration remains a DCC-controlled responsibility."),
    ])

    add_interface_page(doc, 5, "Team and discipline governance", "Invite engineers by discipline, not by individual document", "The team workspace supports project invitations, pending invitation visibility, resend or deletion controls and discipline assignments that prevent engineers from uploading against another discipline's deliverables.", "06-team-resources.png", "Project team and discipline resource controls.", [
        ("Scalable assignment. ", "One discipline invitation can govern a portfolio of planned deliverables."),
        ("Clear onboarding. ", "Pending invitations remain visible until accepted and can be resent or cancelled."),
    ])

    add_interface_page(doc, 6, "Revision record", "Every logical document keeps its full history", "Each MDR entry links to a document-specific record containing metadata, revision history, issue status, file state and secure preview actions.", "15-document-revision-history.png", "Document detail and revision history.", [
        ("One record, many revisions. ", "Teams avoid duplicate document entries while preserving immutable revision evidence."),
        ("Secure access. ", "Preview and download are mediated through authorised routes rather than permanent public links."),
    ])

    add_interface_page(doc, 7, "Submission quality gate", "DCC previews before accepting a revision", "Engineer submissions enter a controlled review queue. The document controller verifies document identity, revision, issue status and file conformity before accepting the revision into the controlled record or returning it with instructions.", "05-submission-review.png", "DCC submission review queue.", [
        ("Quality gate. ", "An uploaded file is not automatically treated as a controlled project issue."),
        ("Visible workload. ", "The navigation count shows how many submissions require DCC action."),
    ])

    add_interface_page(doc, 8, "Actionable notification centre", "Messages explain what happened and where", "Notifications bring together submission, deadline, invitation and review activity. Messages can be opened in place, marked read and cleared without forcing the user into an unrelated page.", "10-notifications.png", "In-app notification centre.", [
        ("Operational context. ", "Submission messages identify discipline, document and revision rather than only stating that a revision exists."),
        ("Reduced noise. ", "Unread counts fall as messages are read and users can clear resolved items."),
    ])

    add_interface_page(doc, 9, "Delivery oversight", "Progress is calculated against the MDR", "The project dashboard compares planned deliverables with uploaded, approved/final and overdue records, then breaks performance down by discipline.", "07-delivery-progress.png", "Project and discipline delivery progress.", [
        ("Plan-based measurement. ", "The denominator is the complete active MDR, not an informal count of files in storage."),
        ("Discipline accountability. ", "Process, piping, structural and electrical performance can be reviewed independently."),
    ])

    add_interface_page(doc, 10, "Management reporting", "A repeatable weekly progress report", "EngiCite produces a client-ready project report using project identity, organisation and client logos, reporting period, overall progress, discipline performance, weekly plan, issued deliverables, variance, look-ahead, challenges and an S-curve.", "14-weekly-project-report.png", "Generated weekly project progress report.", [
        ("Consistent presentation. ", "Project data is transformed into a repeatable report instead of rebuilt in PowerPoint or Excel each week."),
        ("Client-facing logic. ", "Only DCC-accepted records count as issued; internal accept/reject activity is not exposed as client progress."),
    ])

    add_interface_page(doc, 11, "Controlled handover", "Freeze the exact issue set", "The work package workspace lets authorised users assemble accepted revisions by discipline and issue status, then deliver a secure local package or a configured external destination.", "08-work-packages.png", "Engineering work package and delivery workspace.", [
        ("Reproducible package. ", "The selected revision set is frozen so the delivered record cannot silently change later."),
        ("Flexible use. ", "The same controls support interim issue sets, final work packs and formal handover."),
    ])

    add_interface_page(doc, 12, "Client transmission", "Create a professional transmittal from accepted revisions", "The transmittal form captures a unique transmission number, recipient, purpose and cover message, then adds only current DCC-accepted revisions to the issue set.", "09-new-transmittal.png", "New client transmittal form.", [
        ("Controlled selection. ", "Previously issued documents can show the transmittal number, reducing accidental reissue."),
        ("Acknowledgement-ready. ", "The package includes a client-facing transmittal cover suitable for receipt acknowledgement."),
    ])

    add_interface_page(doc, 13, "Evidence retrieval", "Search engineering content, not just filenames", "Evidence search combines document metadata and extracted content so authorised users can locate the right record and navigate back to its source revision.", "11-evidence-search.png", "Engineering evidence search.", [
        ("Faster discovery. ", "Users can search by engineering concepts as well as document numbers and titles."),
        ("Scoped results. ", "Retrieval stays inside the active organisation and project permission boundary."),
    ])

    add_interface_page(doc, 14, "Ask EngiCite", "Project questions must point back to proof", "The AI workspace is designed to answer only from selected, authorised project documents and cite the document number, revision and page. When the evidence is insufficient, the expected response is to say so.", "12-ai-document-chat.png", "Citation-grounded AI document question answering.", [
        ("Evidence first. ", "The differentiator is not a generic chatbot; it is a traceable answer linked to controlled source material."),
        ("Commercial note. ", "Live AI answers require configured OpenAI API billing and production data-control settings."),
    ])

    add_interface_page(doc, 15, "Report governance", "A controlled history of weekly reports", "Reports are stored as a project record, visible by reporting period, overall completion and overdue count. Role permissions determine who can generate or only read the report.", "13-project-reports.png", "Weekly project report register.", [
        ("Auditability. ", "Management can revisit what was reported for a specific week rather than relying on overwritten files."),
        ("Role separation. ", "Project administrators control schedules and generation while other authorised users receive read-only access."),
    ])

    # Capability status
    page_break(doc)
    add_kicker(doc, "Capability status")
    add_title(doc, "Market what is demonstrated - label what still needs configuration", size=22)
    add_deck(doc, "The current build supports a compelling pilot demonstration. External-service and enterprise-readiness claims should remain clearly qualified until the listed controls are configured and evidenced.")
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 3860, 3300])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("Status", "Capability", "Marketing guidance")):
        shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=8.8, color=WHITE, bold=True)
    status_rows = [
        ("DEMONSTRATED", GREEN_LIGHT, GREEN, "Secure sign-in; organisations and projects; role workspaces", "Show in every pilot demo"),
        ("DEMONSTRATED", GREEN_LIGHT, GREEN, "MDR, discipline grouping, Excel import and submission dates", "Lead with DCC time savings"),
        ("DEMONSTRATED", GREEN_LIGHT, GREEN, "Revision upload, history, secure preview and DCC review", "Demonstrate the quality gate"),
        ("DEMONSTRATED", GREEN_LIGHT, GREEN, "Notifications, delivery progress and weekly reports", "Use to sell management visibility"),
        ("DEMONSTRATED", GREEN_LIGHT, GREEN, "Work packages and client transmittal workflow", "Show controlled issue and handover"),
        ("CONFIGURE", ORANGE_LIGHT, AMBER, "AI answers and embeddings", "Requires API credit, processor configuration and production data controls"),
        ("CONFIGURE", ORANGE_LIGHT, AMBER, "Outbound invitation and overdue reminder email", "Requires verified sending domain and Resend configuration"),
        ("CONFIGURE", ORANGE_LIGHT, AMBER, "SharePoint / Google Drive delivery and qualified e-signature", "Sell only after connector and signing-provider acceptance tests"),
        ("RELEASE GATE", "F8E8E6", RED, "Malware scan, pen test, restore exercise, load test and production monitoring", "Do not claim enterprise readiness before evidence is complete"),
    ]
    for label, fill, color, capability, guidance in status_rows:
        cells = table.add_row().cells
        add_status_chip(cells[0], label, fill, color)
        for cell, value in zip(cells[1:], (capability, guidance)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=8.6, color=INK)
    p = doc.add_paragraph(style="Small Note")
    p.paragraph_format.space_before = Pt(6)
    p.add_run("Recommended public phrasing: 'built for secure, role-based engineering document control' until independent security and operational acceptance evidence is available.")

    # Trust and differentiation
    page_break(doc)
    add_kicker(doc, "Trust story and differentiation")
    add_title(doc, "Sell the control model, not only the feature list", size=22)
    add_deck(doc, "The strongest competitive position is the combination of disciplined project delivery and evidence-backed intelligence in a product that remains understandable to engineers, DCCs and managers.")
    add_matrix(doc, ["Trust pillar", "Buyer-facing message", "Proof required"], [
        ("Tenant isolation", "Each organisation and project is separated by policy and server-side authorisation", "Automated two-tenant RLS and IDOR test evidence"),
        ("Least privilege", "Each role sees only the controls needed for its responsibility", "Role matrix, permission tests and denied-action audit events"),
        ("Secure file access", "Files use short-lived authorised access rather than permanent public links", "Signed URL tests, private bucket policies and audit events"),
        ("Traceability", "Uploads, downloads, edits, reviews, transmissions and questions create an audit trail", "Append-only audit verification and export review"),
        ("Responsible AI", "Customer documents are used only to answer authorised project questions, not for model training", "Provider agreement, storage-off configuration and data-flow review"),
        ("Resilience", "Project records can be recovered and exported", "Restore exercise, backup evidence and tenant export/deletion runbook"),
    ], [1900, 4200, 3260], font_size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    add_body(doc, "Competitive distinctions", bold_lead="Competitive distinctions")
    add_bullet_rows(doc, [
        ("Versus shared drives: ", "EngiCite knows the document number, discipline, planned date, revision, acceptance state and transmission history."),
        ("Versus generic DMS platforms: ", "the interface follows engineering delivery language and DCC operating practice from the first screen."),
        ("Versus standalone AI search: ", "answers sit on top of a controlled register and are expected to cite the accepted project source."),
        ("Versus spreadsheet MDRs: ", "the register becomes an active workflow with responsibility, reminders, progress and controlled issue."),
    ], font_size=9.4)

    # Messaging
    page_break(doc)
    add_kicker(doc, "Messaging architecture")
    add_title(doc, "One brand promise, three buyer conversations", size=22)
    add_brand_callout(doc, "Master promise", "CONTROL EVERY DELIVERABLE. CITE EVERY DECISION.")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    add_matrix(doc, ["Audience", "Lead message", "Proof point"], [
        ("Project leadership", "See delivery risk before it becomes a client escalation", "MDR-based progress, overdue work, weekly variance and look-ahead"),
        ("Document control", "Run the register, reviews and transmissions from one controlled workspace", "Excel import, discipline assignments, review queue and frozen issue sets"),
        ("Engineering teams", "Know exactly what is due and submit it without document-control guesswork", "My Deliverables, authorised discipline scope, revision upload and feedback"),
        ("Information security", "Keep project records separated, authorised and auditable", "Tenant policies, signed access, audit events and server-side secrets"),
        ("Knowledge users", "Find the answer and open the proof", "Hybrid retrieval and document/revision/page citations"),
    ], [1750, 4300, 3310], font_size=8.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    add_body(doc, "Approved short-form copy", bold_lead="Approved short-form copy")
    add_bullet_rows(doc, [
        ("Homepage one-liner: ", "Engineering document control, project delivery visibility and evidence-backed answers in one secure workspace."),
        ("30-second pitch: ", "EngiCite turns the MDR into a live delivery system. DCCs organise and control the register, engineers submit against assigned disciplines, managers see progress and risk, and every issued record remains traceable."),
        ("Demo close: ", "From the first planned deliverable to the final client transmittal, EngiCite keeps the responsibility, revision and evidence connected."),
    ], font_size=9.4)

    # Demo and funnel
    page_break(doc)
    add_kicker(doc, "Commercial demonstration")
    add_title(doc, "A 12-minute demo that follows one deliverable", size=22)
    add_deck(doc, "Avoid touring every menu. Tell one end-to-end story that makes the role separation and traceability obvious.")
    add_matrix(doc, ["Time", "Show", "Commercial point"], [
        ("0:00-1:00", "Landing page and organisation workspace", "Clear category promise and tenant separation"),
        ("1:00-3:00", "Import or open the MDR grouped by discipline", "Fast mobilisation and a controlled plan"),
        ("3:00-4:30", "Invite a discipline engineer and show assigned work", "Ownership without exposing unrelated controls"),
        ("4:30-6:00", "Open document history and secure preview", "Immutable revision evidence and authorised access"),
        ("6:00-7:30", "DCC review and acceptance", "Quality gate before a document becomes controlled"),
        ("7:30-9:00", "Progress dashboard and weekly report", "Early warning and repeatable management reporting"),
        ("9:00-10:30", "Create a transmittal or work package", "Controlled issue set and professional client handover"),
        ("10:30-12:00", "Evidence search and Ask EngiCite", "Open the proof behind a project answer"),
    ], [1200, 3600, 4560], font_size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_brand_callout(doc, "Recommended pilot close", "LET US LOAD ONE REAL PROJECT MDR, ASSIGN TWO DISCIPLINES AND RUN ONE DOCUMENT FROM PLANNED DATE TO CLIENT TRANSMITTAL.")

    # GTM plan
    page_break(doc)
    add_kicker(doc, "90-day go-to-market plan")
    add_title(doc, "Lead with a controlled pilot, then convert proof into repeatable demand", size=22)
    add_deck(doc, "The first objective is not broad awareness. It is credible product proof with a small number of high-fit engineering organisations and document-control champions.")
    add_matrix(doc, ["Phase", "Commercial objective", "Actions", "Exit signal"], [
        ("Days 1-30 - Prepare", "Make the product safe and demo-ready", "Complete deployment, email/AI configuration, demo dataset, security evidence pack, pilot agreement and demo script", "A repeatable 12-minute demo and a signed pilot checklist"),
        ("Days 31-60 - Prove", "Run 2-3 design-partner pilots", "Import real MDRs, onboard DCC and engineers, measure setup time, overdue visibility and transmittal effort", "At least one complete deliverable workflow and quantified customer outcome"),
        ("Days 61-90 - Convert", "Turn outcomes into demand", "Publish a controlled case study, webinar, DCC guide, comparison page and partner outreach campaign", "Qualified pipeline, reference customer and paid conversion proposal"),
    ], [1500, 2300, 3650, 1910], font_size=8.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_body(doc, "Priority channels", bold_lead="Priority channels")
    add_bullet_rows(doc, [
        ("Founder-led outreach: ", "project directors, lead document controllers and engineering managers in existing professional networks."),
        ("Design-partner workshops: ", "60-minute MDR health assessment using the customer's current spreadsheet and workflow."),
        ("LinkedIn authority content: ", "practical DCC topics - late deliverables, revision errors, transmittal control and handover readiness."),
        ("Implementation partners: ", "document-control consultants, EPC digitalisation specialists and engineering project-service firms."),
    ], font_size=9.2)

    # Content, funnel and packaging
    page_break(doc)
    add_kicker(doc, "Campaign, funnel and packaging")
    add_title(doc, "Convert operational pain into a measurable pilot", size=22)
    add_matrix(doc, ["Funnel stage", "Buyer question", "Content / offer", "Success metric"], [
        ("Awareness", "Why do our MDR and transmittals still consume so much time?", "Document Control Maturity Checklist; short workflow videos", "Relevant engagement and checklist downloads"),
        ("Consideration", "Can EngiCite fit our project and roles?", "Role-based product tour; MDR import demonstration; security overview", "Qualified discovery meetings"),
        ("Evaluation", "Will it work with our real records?", "One-project controlled pilot with agreed success criteria", "Pilot activation and weekly active users"),
        ("Conversion", "Can we deploy safely and justify cost?", "Outcome report, implementation plan, security evidence and commercial proposal", "Paid project or annual subscription"),
        ("Expansion", "Can we standardise across projects?", "Portfolio dashboard, reusable templates and champion enablement", "Additional projects, seats and storage"),
    ], [1500, 2750, 3300, 1810], font_size=8.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)
    add_body(doc, "Packaging hypothesis - validate through pilots", bold_lead="Packaging hypothesis - validate through pilots")
    add_matrix(doc, ["Package", "Designed for", "Commercial boundary"], [
        ("Pilot", "One project team proving the workflow", "Limited term, project, seats and storage; guided onboarding; AI optional"),
        ("Project", "A consultancy or contractor running active projects", "Multiple roles, MDR, revisions, reporting, transmittals and defined usage limits"),
        ("Enterprise", "Owner-operators and multi-project organisations", "Portfolio controls, SSO/SCIM, private connectivity/region options, advanced retention and integrations"),
    ], [1650, 3550, 4160], font_size=8.8)
    p = doc.add_paragraph(style="Small Note")
    p.paragraph_format.space_before = Pt(6)
    p.add_run("Do not publish fixed prices until at least three pilot discovery cycles establish seat counts, storage, implementation effort and willingness to pay. Price the controlled project outcome, not the number of documents.")

    # Measurement and next actions
    page_break(doc)
    add_kicker(doc, "Measurement and next decisions")
    add_title(doc, "Prove adoption, control quality and commercial value", size=22)
    add_deck(doc, "A successful launch should show that teams complete controlled work faster and with fewer avoidable errors - not merely that users signed in.")
    add_matrix(doc, ["Measure", "Pilot definition", "Why it matters"], [
        ("Time to first controlled project", "Days from agreement to MDR imported, roles assigned and first revision accepted", "Tests onboarding friction and implementation cost"),
        ("DCC cycle time", "Submission to accepted/returned decision", "Measures operational responsiveness"),
        ("On-time submission rate", "Due deliverables submitted by planned date", "Connects EngiCite to project performance"),
        ("Revision rework rate", "Submissions returned for metadata/file nonconformance", "Measures quality improvement"),
        ("Transmittal preparation time", "Minutes from selection to completed issue package", "Quantifies direct DCC time saving"),
        ("Evidence retrieval time", "Time to find the correct revision/page for a question", "Quantifies search and AI value"),
        ("Pilot-to-paid conversion", "Pilots converting to a project or annual plan", "Tests commercial fit"),
    ], [2100, 3900, 3360], font_size=8.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)
    add_brand_callout(doc, "Immediate next decision", "SELECT ONE REPRESENTATIVE PROJECT AND DEFINE A PILOT SCORECARD BEFORE ADDING MORE FEATURE SCOPE.")
    add_bullet_rows(doc, [
        ("1. ", "Prepare a clean demonstration organisation with approved fictitious project, client and document data."),
        ("2. ", "Complete production email, AI, processing, monitoring, backup and security evidence gates."),
        ("3. ", "Recruit two DCC-led design partners and agree measurable before/after outcomes."),
        ("4. ", "Record the 12-minute demo and produce a one-page pilot offer from this blueprint."),
    ], font_size=9.2)

    # Final page
    page_break(doc)
    for _ in range(4):
        doc.add_paragraph()
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing_logo = p_logo.add_run().add_picture(str(LOGO), width=Inches(4.4))._inline
    closing_logo.docPr.set("descr", "EngiCite company logo")
    add_title(doc, "Know the answer. Cite the proof.", size=28, color=NAVY, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_deck(doc, "EngiCite connects the engineering delivery plan, controlled document record and project evidence into one secure workflow.", align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("CONTROL EVERY DELIVERABLE  |  CITE EVERY DECISION"), size=10, color=ORANGE, bold=True)
    p = doc.add_paragraph(style="Small Note")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    p.add_run("Internal blueprint. Replace demonstration project and customer data with approved marketing data before public distribution.")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
