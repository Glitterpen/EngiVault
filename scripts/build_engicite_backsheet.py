from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('docs/templates/EngiCite_Document_Backsheet.docx')
LOGO = Path('promotional/engicite-logo-word.png')
NAVY = '10243E'; ORANGE = 'E8733F'; MUTED = '617083'; PALE = 'F4F6F8'; LINE = 'DCE2E9'

def set_font(run, size, color=NAVY, bold=False, italic=False):
    run.font.name = 'Arial'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold; run.italic = italic

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def cell_margins(cell, top=180, start=260, bottom=180, end=260):
    tcPr = cell._tc.get_or_add_tcPr(); mar = OxmlElement('w:tcMar')
    for name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement('w:' + name); node.set(qn('w:w'), str(value)); node.set(qn('w:type'), 'dxa'); mar.append(node)
    tcPr.append(mar)

doc = Document(); section = doc.sections[0]
section.page_width = Inches(8.5); section.page_height = Inches(11)
section.top_margin = Inches(.55); section.bottom_margin = Inches(.48)
section.left_margin = Inches(.62); section.right_margin = Inches(.62)
section.header_distance = Inches(.25); section.footer_distance = Inches(.25)

normal = doc.styles['Normal']; normal.font.name = 'Arial'; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(0)

# Brand artwork hero.
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(20)
p.add_run().add_picture(str(LOGO), width=Inches(7.26), height=Inches(4.085))

# Small orange rule.
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(18)
r = p.add_run('━━━━━━'); set_font(r, 10, ORANGE, True)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(7)
set_font(p.add_run('ENGINEERING INTELLIGENCE,'), 17, NAVY, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(13)
set_font(p.add_run('EVIDENCED.'), 24, ORANGE, True)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(24)
set_font(p.add_run('Controlled documents. Verifiable answers. Defensible decisions.'), 11.5, MUTED, False, True)

# Contact panel.
table = doc.add_table(rows=1, cols=1); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
table.columns[0].width = Inches(6.55)
cell = table.cell(0, 0); cell.width = Inches(6.55); shade(cell, PALE); cell_margins(cell, 220, 260, 220, 260); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
cell.text = ''
p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(7)
set_font(p.add_run('KNOW THE ANSWER.  CITE THE PROOF.'), 10.5, NAVY, True)
p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
set_font(p.add_run('www.engicite.com'), 10, ORANGE, True)
p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('hello@engicite.com'), 10, MUTED, False)

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(5); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('ENGICITE'), 9, NAVY, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3)
set_font(p.add_run('Engineering knowledge, with the proof attached.'), 8.5, MUTED, False)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('© 2026 EngiCite. All rights reserved.'), 8, MUTED, False)

# Quiet confidentiality line at the foot of the backsheet.
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run('CONFIDENTIAL  •  For the intended recipient only'), 7.5, MUTED, True)

doc.core_properties.title = 'EngiCite Document Backsheet'
doc.core_properties.subject = 'Reusable branded back cover for EngiCite documents'
doc.core_properties.author = 'EngiCite'
OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT)
print(OUT.resolve())
