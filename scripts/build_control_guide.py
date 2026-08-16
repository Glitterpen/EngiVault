from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "EngiCite_Levels_of_Control_Guide.docx"
LOGO = ROOT / "apps" / "web" / "public" / "engicite-logo-transparent.png"
SKILL_SCRIPTS = Path(
    r"C:\Users\ahmed\.codex\plugins\cache\openai-primary-runtime\documents\26.805.11740\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


NAVY = "071A36"
INK = "24384F"
MUTED = "617083"
ORANGE = "F97316"
GREEN = "0C5B45"
PALE_GREEN = "E8F1ED"
PALE_ORANGE = "FFF2EB"
PALE_BLUE = "EEF3F7"
WHITE = "FFFFFF"
BORDER = "DCE5EA"
LIGHT_BORDER = "E9EEF1"
LETTER_WIDTH_DXA = 9360


def set_run(run, *, size=None, color=INK, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    tr_pr.append(cannot_split)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def add_page_field(paragraph, field="PAGE"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run(run, size=8.5, color=MUTED)


def add_numbering_definition(doc, *, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def add_bullet(doc, text, bullet_num, *, bold_lead=None):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_num)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), size=10.5, bold=True, color=NAVY)
        set_run(p.add_run(text[len(bold_lead):]), size=10.5, color=INK)
    else:
        set_run(p.add_run(text), size=10.5, color=INK)
    return p


def add_numbered_step(doc, title, detail, number_num):
    p = doc.add_paragraph()
    apply_numbering(p, number_num)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(title), size=11, bold=True, color=NAVY)
    detail_p = doc.add_paragraph()
    detail_p.paragraph_format.left_indent = Inches(0.375)
    detail_p.paragraph_format.space_after = Pt(8)
    detail_p.paragraph_format.line_spacing = 1.15
    set_run(detail_p.add_run(detail), size=10.2, color=INK)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Title", 30, 0, 8, NAVY),
        ("Subtitle", 13.5, 0, 18, MUTED),
        ("Heading 1", 16, 18, 10, NAVY),
        ("Heading 2", 13, 14, 7, GREEN),
        ("Heading 3", 12, 10, 5, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name not in ("Subtitle",)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def configure_running_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    if LOGO.exists():
        logo = p.add_run().add_picture(str(LOGO), width=Inches(1.25))
        logo._inline.docPr.set("descr", "EngiCite corporate logo")
        logo._inline.docPr.set("title", "EngiCite")
    right = p.add_run("   CONTROL & RESPONSIBILITY GUIDE")
    set_run(right, size=8, bold=True, color=MUTED)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    left, right_cell = table.rows[0].cells
    left_p = left.paragraphs[0]
    left_p.paragraph_format.space_after = Pt(0)
    set_run(left_p.add_run("EngiCite | Product operating guide"), size=8.5, color=MUTED)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_after = Pt(0)
    set_run(right_p.add_run("Page "), size=8.5, color=MUTED)
    add_page_field(right_p)
    apply_table_geometry(table, [7000, 2360], indent_dxa=120, cell_margins_dxa={"top": 0, "bottom": 0, "start": 120, "end": 120})
    for cell in table.rows[0].cells:
        set_cell_borders(cell, color=WHITE, size="0")


def add_callout(doc, label, text, *, fill=PALE_GREEN, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color=fill, size="0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(label.upper()), size=8.5, bold=True, color=accent)
    body = cell.add_paragraph()
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = 1.15
    set_run(body.add_run(text), size=10.5, color=INK)
    apply_table_geometry(table, [9360], indent_dxa=180, cell_margins_dxa={"top": 130, "bottom": 130, "start": 180, "end": 180})
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(46)
    logo = p.add_run().add_picture(str(LOGO), width=Inches(4.4))
    logo._inline.docPr.set("descr", "EngiCite corporate logo")
    logo._inline.docPr.set("title", "EngiCite")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    set_run(kicker.add_run("GOVERNANCE & OPERATING GUIDE"), size=9, bold=True, color=ORANGE)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run(title.add_run("Levels of Control"), size=30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        subtitle.add_run("Who controls organisations, projects, engineering documents, approvals and client transmission"),
        size=13.5,
        color=MUTED,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    add_callout(
        doc,
        "Core governance principle",
        "EngiCite separates management, document control, engineering production and read-only access. No individual should create, submit, approve and transmit the same controlled document without independent oversight.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )

    meta = doc.add_table(rows=4, cols=2)
    values = [
        ("Document owner", "EngiCite Product Team"),
        ("Version", "1.0"),
        ("Date", "12 August 2026"),
        ("Audience", "Administrators, project managers, DCC, engineers and viewers"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        set_cell_shading(row.cells[0], PALE_BLUE)
        set_cell_shading(row.cells[1], WHITE)
        for cell in row.cells:
            set_cell_borders(cell, LIGHT_BORDER, "5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_run(row.cells[0].paragraphs[0].add_run(label), size=9, bold=True, color=NAVY)
        set_run(row.cells[1].paragraphs[0].add_run(value), size=9.5, color=INK)
    apply_table_geometry(meta, [2100, 7260], indent_dxa=120)
    doc.add_page_break()


def add_role_matrix(doc):
    doc.add_heading("1. Control architecture", level=1)
    p = doc.add_paragraph()
    set_run(
        p.add_run("The platform applies role separation at three practical levels: company governance, project execution and controlled document delivery."),
        size=10.8,
        color=INK,
    )
    add_callout(
        doc,
        "Operating rule",
        "Access is granted only to the organisation, project, discipline and action required by the user's assigned role. Higher visibility does not automatically mean permission to perform DCC or engineering tasks.",
    )

    doc.add_heading("Role and authority matrix", level=2)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Role", "Primary control", "Permitted activities", "Important restriction"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_borders(cell, NAVY, "6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=8.7, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])

    rows = [
        (
            "EngiCite Super Administrator",
            "The SaaS platform and its availability.",
            "Manages subscriptions, platform health, technical support and controlled system administration.",
            "Has no routine access to customer documents. Exceptional support access must be authorised and audited.",
        ),
        (
            "Organisation Administrator",
            "The company workspace and portfolio.",
            "Creates organisations and projects, appoints project leadership, allocates resources and monitors project health, progress and issues.",
            "Does not normally create MDR entries, upload engineering revisions, approve submissions or issue transmittals.",
        ),
        (
            "Project Administrator / Project Manager",
            "One project's objectives, team and delivery performance.",
            "Maintains the project brief, objectives, team, disciplines, schedule, resources and overall progress.",
            "Does not replace the Document Controller's formal document-control authority.",
        ),
        (
            "Document Controller (DCC)",
            "The official MDR, revision acceptance and client issue record.",
            "Creates or imports the MDR, assigns disciplines, invites engineers, sets dates, previews submissions, accepts or rejects revisions and issues controlled transmittals.",
            "Does not upload revisions as the originating engineer and cannot bypass acceptance requirements.",
        ),
        (
            "Discipline Engineer",
            "Production and submission of assigned engineering deliverables.",
            "Views the project brief and assigned discipline, then uploads revisions against allocated MDR documents and deadlines.",
            "Cannot create the MDR, upload into another discipline, approve their own work or issue client transmittals.",
        ),
        (
            "Viewer",
            "Read-only access for authorised stakeholders.",
            "Views permitted projects, documents, revisions, progress and approved records.",
            "Cannot create, edit, upload, approve, invite, delete or transmit controlled information.",
        ),
    ]
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for col_idx, text in enumerate(values):
            cell = cells[col_idx]
            set_cell_shading(cell, "F8FAFB" if row_idx % 2 else WHITE)
            set_cell_borders(cell, LIGHT_BORDER, "5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_run(p.add_run(text), size=8.4 if col_idx else 8.7, bold=(col_idx == 0), color=NAVY if col_idx == 0 else INK)

    apply_table_geometry(table, [1700, 2000, 3000, 2660], indent_dxa=120, cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120})


def add_workflow(doc, number_num):
    doc.add_heading("2. Controlled document workflow", level=1)
    intro = doc.add_paragraph()
    set_run(intro.add_run("The intended control chain is sequential. Each stage hands responsibility to a different authorised role."), size=10.8)
    steps = [
        ("Organisation and project established", "The Organisation Administrator creates the workspace and project, then appoints project leadership and resources."),
        ("Project context and team defined", "The Project Manager records the introduction, objectives, disciplines, schedule and responsible team."),
        ("Master Document Register controlled", "The DCC creates or imports planned deliverables, assigns disciplines and records agreed submission dates."),
        ("Revision produced and submitted", "The assigned Discipline Engineer uploads the correct revision only against an authorised document in their discipline."),
        ("Submission independently checked", "The DCC previews the file for conformance, then accepts it or rejects it with a reason and notification."),
        ("Accepted records issued to the client", "The DCC selects accepted, processing-ready revisions and freezes them into a numbered transmittal with acknowledgement evidence."),
    ]
    for title, detail in steps:
        add_numbered_step(doc, title, detail, number_num)
    add_callout(
        doc,
        "Approval boundary",
        "An engineer cannot approve their own submission. Only a DCC-accepted revision can be included in a controlled client transmittal.",
        fill=PALE_ORANGE,
        accent=ORANGE,
    )


def add_security_controls(doc, bullet_num):
    doc.add_heading("3. Access and security boundaries", level=1)
    controls = [
        ("Organisation isolation: ", "Users in one organisation cannot see another organisation's information."),
        ("Project isolation: ", "Access to one project does not automatically grant access to other projects."),
        ("Discipline isolation: ", "An engineer can submit only to documents belonging to their assigned discipline."),
        ("Revision control: ", "Each file stays linked to its exact document number, revision and issue status."),
        ("Approval control: ", "Submission and acceptance are performed by separate authorised roles."),
        ("Transmission control: ", "Only accepted, ready revisions can enter a numbered client transmittal."),
        ("File-access control: ", "Downloads are provided through temporary signed links instead of public file locations."),
        ("Audit control: ", "Uploads, reviews, downloads, invitations, AI questions and transmissions are recorded."),
    ]
    for lead, detail in controls:
        add_bullet(doc, lead + detail, bullet_num, bold_lead=lead)

    doc.add_heading("4. Quick responsibility guide", level=1)
    quick = [
        ("Create an organisation or project", "Organisation Administrator"),
        ("Define project objectives, team and schedule", "Project Manager"),
        ("Create or import the MDR", "Document Controller"),
        ("Upload an engineering revision", "Assigned Discipline Engineer"),
        ("Preview, accept or reject a submission", "Document Controller"),
        ("Prepare and issue a client transmittal", "Document Controller"),
        ("View approved information without changing it", "Viewer"),
    ]
    table = doc.add_table(rows=1, cols=2)
    for idx, text in enumerate(("Required action", "Authorised role")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN)
        set_cell_borders(cell, GREEN, "6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=9, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for idx, (action, role) in enumerate(quick):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for cell in cells:
            set_cell_shading(cell, WHITE if idx % 2 == 0 else "F6F9F7")
            set_cell_borders(cell, LIGHT_BORDER, "5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_run(cells[0].paragraphs[0].add_run(action), size=9.4, color=INK)
        set_run(cells[1].paragraphs[0].add_run(role), size=9.4, bold=True, color=GREEN)
    apply_table_geometry(table, [6300, 3060], indent_dxa=140, cell_margins_dxa={"top": 90, "bottom": 90, "start": 140, "end": 140})


def add_final_notes(doc, bullet_num):
    doc.add_heading("5. Control outcome", level=1)
    p = doc.add_paragraph()
    set_run(
        p.add_run("When these controls are applied together, EngiCite provides a defensible chain of custody from project setup to client issue."),
        size=10.8,
        color=INK,
    )
    for text in (
        "Management receives portfolio visibility without taking over specialist document-control actions.",
        "Engineers receive clear assignments and upload rights limited to their disciplines.",
        "The DCC maintains an independent official record and controls what is formally issued.",
        "Clients receive traceable document packages linked to numbered transmittals and acknowledgement evidence.",
    ):
        add_bullet(doc, text, bullet_num)
    add_callout(
        doc,
        "Implementation note",
        "This guide describes the intended EngiCite operating model. Before production deployment, each role and boundary should be verified against the live permission tests, database policies and audit records.",
        fill=PALE_BLUE,
        accent=NAVY,
    )


def build():
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    configure_running_header_footer(doc.sections[0])
    bullet_num = add_numbering_definition(doc, bullet=True)
    number_num = add_numbering_definition(doc, bullet=False)

    add_cover(doc)
    add_role_matrix(doc)
    add_workflow(doc, number_num)
    add_security_controls(doc, bullet_num)
    add_final_notes(doc, bullet_num)

    core = doc.core_properties
    core.title = "EngiCite Levels of Control and Responsibility Guide"
    core.subject = "Role separation, controlled document workflow and access boundaries"
    core.author = "EngiCite Product Team"
    core.keywords = "EngiCite, document control, roles, permissions, MDR, transmittal"
    core.comments = "Generated as an EngiCite product operating guide."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
