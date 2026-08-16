from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('docs/legal/EngiCite_Non_Disclosure_Agreement.docx')
NAVY = '10243E'; ORANGE = 'E8733F'; GRAY = '617083'; LIGHT = 'EEF2F6'; BORDER = 'D7DEE7'

def font(run, size=10.5, bold=False, color='111111', italic=False):
    run.font.name = 'Arial'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Arial'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd')) or OxmlElement('w:shd'); shd.set(qn('w:fill'), fill)
    if shd.getparent() is None: tcPr.append(shd)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar') or OxmlElement('w:tcMar')
    if tcMar.getparent() is None: tcPr.append(tcMar)
    for tag, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el = tcMar.find(qn('w:'+tag)) or OxmlElement('w:'+tag); el.set(qn('w:w'), str(val)); el.set(qn('w:type'),'dxa')
        if el.getparent() is None: tcMar.append(el)

def set_cell_text(cell, text, bold=False, color='111111', size=9.5):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); font(r,size,bold,color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def add_clause(doc, heading, paras):
    doc.add_heading(heading, level=1)
    for item in paras:
        if isinstance(item, tuple):
            label, text = item; p=doc.add_paragraph(); p.style='List Bullet'; r=p.add_run(label+' '); font(r,bold=True); font(p.add_run(text))
        else:
            p=doc.add_paragraph(item)

def signature_block(doc, title):
    doc.add_heading(title, level=2)
    rows=[('Legal name',''),('Signature',''),('Name and title',''),('Date',''),('Email','')]
    t=doc.add_table(rows=0, cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    for label,value in rows:
        c=t.add_row().cells; c[0].width=Inches(1.55); c[1].width=Inches(4.95); set_cell_text(c[0],label,True,NAVY); set_cell_text(c[1],value+'\n')

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.78); sec.bottom_margin=Inches(.72); sec.left_margin=Inches(.88); sec.right_margin=Inches(.88); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string('111111'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.12
for name,size,before,after,color in [('Heading 1',15,14,6,NAVY),('Heading 2',12,10,4,NAVY),('Heading 3',10.5,8,3,GRAY)]:
    s=styles[name]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for name in ['List Bullet','List Number']:
    s=styles[name]; s.font.name='Arial'; s.font.size=Pt(10.5); s.paragraph_format.left_indent=Inches(.28); s.paragraph_format.first_line_indent=Inches(-.18); s.paragraph_format.space_after=Pt(4)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run('ENGICITE  |  CONFIDENTIAL'),8,True,GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run('Non-Disclosure Agreement  •  '),8,False,GRAY)
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(4); font(p.add_run('ENGICITE'),12,True,ORANGE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); font(p.add_run('NON-DISCLOSURE'),28,True,NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18); font(p.add_run('CONFIDENTIALITY, INTELLECTUAL PROPERTY AND SECURITY AGREEMENT'),16,True,NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(16); font(p.add_run('For employees, contractors, advisers, secondees, interns and other authorised team members'),10.5,False,GRAY,True)

t=doc.add_table(rows=0,cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
meta=[('Effective Date','[DD MONTH YYYY]'),('Company','[FULL LEGAL NAME OF ENGICITE OPERATOR]'),('Registered Address','[REGISTERED ADDRESS]'),('Team Member / Service Provider','[FULL LEGAL NAME]'),('Engagement / Team','[ROLE, DEPARTMENT OR PROJECT]'),('Governing Law and Forum','[JURISDICTION AND COURTS / ARBITRATION SEAT]')]
for i,(a,b) in enumerate(meta):
    c=t.add_row().cells; c[0].width=Inches(2.05); c[1].width=Inches(4.55); shade(c[0],LIGHT); set_cell_text(c[0],a,True,NAVY); set_cell_text(c[1],b)

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(0); font(p.add_run('IMPORTANT COMPLETION NOTE'),9,True,ORANGE)
p=doc.add_paragraph('Complete every bracketed field before signature. If the team member is supplied through a service provider, both the provider and each individual with access should sign or execute an approved joinder.'); p.paragraph_format.space_after=Pt(12)

doc.add_heading('Agreement',level=1)
doc.add_paragraph('This Agreement is made on the Effective Date between the Company and the Team Member identified above. If the Team Member is an entity, it enters this Agreement for itself and must ensure that every representative receiving Confidential Information is bound in writing by obligations at least as protective as these. The Company and the Team Member are each a “Party” and together the “Parties”.')
doc.add_paragraph('The Team Member may receive access to EngiCite, a secure multi-tenant platform for controlled engineering documents, master document registers, revision management, retrieval, comparison and citation-grounded AI. Access may expose highly sensitive customer, technical, operational and commercial information. In consideration of the engagement, access, training and exchange of information, the Parties agree as follows.')

add_clause(doc,'1. Purpose and authorised use',[
    'The Team Member may use Confidential Information solely to perform duties expressly authorised by the Company in connection with the Engagement (the “Purpose”). The Team Member must follow the Company’s current policies, instructions, access controls and customer-specific restrictions.',
    ('Need-to-know.','Access, copy, discuss and share Confidential Information only where genuinely required for the Purpose and only with Authorised Persons.'),
    ('No personal benefit.','Do not use Confidential Information to compete, trade, solicit business, train an external system, build a personal portfolio, publish research, or benefit any third party.'),
])

add_clause(doc,'2. Definitions',[
    ('“Authorised Person” means','the Team Member and any representative whom the Company has expressly authorised, who needs the information for the Purpose, and who is bound by written confidentiality and security duties at least as protective as this Agreement.'),
    ('“Confidential Information” means','all non-public information disclosed or made accessible before or after the Effective Date, in any form, that is marked confidential or that a reasonable person would understand to be confidential given its nature or the circumstances.'),
    ('“Customer Content” means','documents, drawings, models, spreadsheets, metadata, master document registers, revisions, prompts, questions, answers, citations, extracted text, embeddings, access records and other information supplied by or processed for a customer, prospect, partner or user.'),
    ('“Company Systems” means','EngiCite environments, repositories, source control, databases, storage, AI/retrieval pipelines, devices, networks, credentials, logs and connected services.'),
    ('“Work Product” means','all designs, discoveries, documentation, code, configurations, prompts, workflows, models, tests, inventions, analyses and other materials created specifically for the Company in the course of the Engagement.'),
])

add_clause(doc,'3. Information covered',[
    'Confidential Information includes, without limitation:',
    ('Product and technology.','Source code, architecture, schemas, APIs, algorithms, prompts, evaluation methods, retrieval and ranking logic, model/provider configurations, roadmaps, prototypes, vulnerabilities, test results and deployment details.'),
    ('Customer and project information.','Customer identities, contracts, pricing, engineering documents, project data, document numbers, revisions, drawings, specifications, asset information, user activity and Customer Content.'),
    ('Security and operations.','Credentials, keys, tokens, access-control rules, tenant identifiers, incident information, audit data, backups, retention settings, threat models and business-continuity materials.'),
    ('Business and people.','Financial data, forecasts, investor information, strategy, sales pipeline, supplier terms, personnel information, compensation, recruitment and internal communications.'),
    ('Third-party information.','Information the Company must protect for customers, licensors, partners, employees and other third parties.'),
])

add_clause(doc,'4. Exclusions and burden of proof',[
    'Confidential Information does not include information the Team Member can prove with contemporaneous written records: (a) was lawfully known without restriction before disclosure; (b) becomes public through no breach of duty; (c) is received lawfully from a third party without confidentiality duty; or (d) is independently developed without using or referring to Confidential Information. A combination of public elements is not excluded if the combination itself is non-public. The Team Member bears the burden of establishing an exclusion.'
])

add_clause(doc,'5. Confidentiality duties',[
    ('Protect.','Use at least reasonable care and no less care than used for similarly sensitive personal information; preserve confidentiality, integrity and availability.'),
    ('Limit disclosure.','Do not disclose outside Authorised Persons without prior written approval from the Company’s designated legal or security contact.'),
    ('Minimise.','Access and retain only the minimum information required, avoid unnecessary downloads or copies, and use approved redacted or synthetic data where practicable.'),
    ('No circumvention.','Do not bypass tenant isolation, authorisation, logging, usage limits, content restrictions, or technical safeguards, even for testing, unless specifically authorised in writing.'),
    ('Responsibility.','Remain responsible for breaches by the Team Member’s representatives and immediately stop unauthorised access or use.'),
])

add_clause(doc,'6. AI, data and engineering-specific restrictions',[
    ('Approved AI only.','Do not submit Confidential Information, Customer Content, source code, credentials or internal prompts to public, personal or unapproved generative-AI, transcription, translation, code-assistant or analytics services.'),
    ('No training or benchmarking.','Do not use Confidential Information or Customer Content to train, fine-tune, evaluate, benchmark or improve any model or dataset except within a Company-approved workflow and documented customer/provider terms.'),
    ('Grounded outputs.','Treat AI outputs, summaries, comparisons and citations as potentially incomplete or incorrect. Verify material outputs against authorised source documents before operational, engineering, safety, commercial or customer use.'),
    ('Tenant isolation.','Never mix, compare or expose content across organisations or projects unless expressly authorised. Do not use one customer’s content to answer another customer’s question.'),
    ('No safety reliance.','Do not represent EngiCite output as a substitute for professional engineering judgment, controlled-document review, formal approval, or applicable safety and regulatory processes.'),
    ('Export controls and sensitivity.','Observe customer classifications, data-residency requirements, export controls and handling restrictions communicated by the Company.'),
])

add_clause(doc,'7. Security requirements',[
    ('Accounts.','Use only assigned accounts; never share credentials; use multi-factor authentication where provided; maintain strong unique passwords through an approved password manager.'),
    ('Devices and storage.','Use approved, patched, encrypted devices and Company-approved storage. Do not store Confidential Information on personal email, consumer cloud drives, removable media or unapproved messaging apps.'),
    ('Code and secrets.','Keep secrets out of source code, tickets, screenshots and logs; use approved secret stores; do not copy Company code to personal repositories or devices.'),
    ('Working practices.','Lock screens, protect conversations and displays, avoid public networks unless protected, follow clean-desk practices where appropriate, and securely dispose of materials.'),
    ('Monitoring.','To the extent permitted by law, Company Systems may log and monitor access and activity for security, compliance, support and audit purposes; this does not authorise access beyond the Purpose.'),
])

add_clause(doc,'8. Security and confidentiality incidents',[
    'The Team Member must report any actual or suspected loss, unauthorised access, disclosure, malware, phishing, credential exposure, tenant crossover, misdirected communication, vulnerable configuration or policy breach immediately, and in all cases within [FOUR (4)] hours after discovery, to [SECURITY EMAIL / HOTLINE].',
    'The Team Member must preserve relevant evidence, follow containment instructions, cooperate fully with investigation and remediation, and must not notify customers, regulators, the media or other third parties unless authorised in writing or legally required. Reporting does not replace emergency action needed to prevent imminent harm.'
])

add_clause(doc,'9. Compelled disclosure and protected reporting',[
    'If legally compelled to disclose Confidential Information, the Team Member must, where lawful, promptly notify the Company in writing, disclose only the minimum legally required, seek confidential treatment, and reasonably assist the Company in seeking protection. Nothing in this Agreement prohibits protected whistleblowing, reporting suspected unlawful conduct to a competent authority, cooperating with regulators, or making disclosures that cannot lawfully be restricted.'
])

add_clause(doc,'10. Ownership; no licence',[
    'All Confidential Information remains the property of the Company or its applicable customer/licensor. No licence or other right is granted by disclosure except the limited, revocable right to use it for the Purpose. The Team Member must not remove confidentiality, copyright, ownership or document-control notices. Feedback regarding EngiCite may be used by the Company without restriction or payment, provided it does not identify the Team Member’s unrelated confidential information.'
])

add_clause(doc,'11. Work Product and intellectual property',[
    'To the fullest extent permitted by law, Work Product created within the scope of the Engagement is specially commissioned for the Company. The Team Member hereby assigns to the Company, with full title guarantee, all worldwide rights, title and interest in Work Product, including intellectual-property rights, upon creation. The Team Member will sign documents and provide reasonable assistance needed to confirm, register or enforce those rights.',
    'The Team Member retains ownership of tools, materials and intellectual property developed independently before the Engagement and identified in Schedule 2 (“Background Materials”). If Background Materials are embedded in Work Product, the Team Member grants the Company and its successors a perpetual, worldwide, irrevocable, transferable, sublicensable, royalty-free licence to use, reproduce, modify, distribute, commercialise and create derivative works from them as part of or in connection with the Work Product.',
    'The Team Member must not incorporate third-party code, datasets, content or open-source software into Work Product except in accordance with Company approval and licence-compliance procedures. To the extent permitted by law, the Team Member waives and agrees not to assert moral rights in Work Product.'
])

add_clause(doc,'12. Personal data and privacy',[
    'The Team Member may process personal data only on documented Company instructions, for the Purpose, and in accordance with applicable privacy law and Company policy. The Team Member must not create unauthorised datasets, export personal data, attempt re-identification, or use personal data for profiling or unrelated analytics. Privacy requests, complaints and regulatory contacts must be promptly forwarded to [PRIVACY CONTACT].'
])

add_clause(doc,'13. Return, deletion and access termination',[
    'Upon request or the end of the Engagement, the Team Member must immediately stop use, return Company property, surrender credentials, and within [FIVE (5)] business days securely delete all Confidential Information in the Team Member’s possession or control, including local copies and exports. The Team Member must certify completion if requested. Automatic backups may be retained only where deletion is impracticable, access is restricted, normal-cycle deletion applies, and this Agreement continues to protect them. The Company may terminate or suspend access at any time.'
])

add_clause(doc,'14. Term and survival',[
    'This Agreement begins on the Effective Date and continues throughout the Engagement. Confidentiality and restricted-use duties continue for five (5) years after the later of the last disclosure or termination, except that duties for trade secrets, credentials, security vulnerabilities, personal data and Customer Content continue for as long as the information remains protected or confidential under applicable law, customer obligation or its nature. Clauses on ownership, Work Product, return/deletion, remedies, governing law and any provisions intended by their nature to survive will survive termination.'
])

add_clause(doc,'15. Representations and relationship',[
    'The Team Member represents that entering and performing this Agreement does not breach another duty and will not bring or use another person’s confidential information without lawful authorisation. Confidential Information is provided “as is” without warranty except as expressly agreed in writing. This Agreement does not require disclosure, guarantee continued engagement, create a partnership or agency, or authorise either Party to bind the other.'
])

add_clause(doc,'16. Remedies and liability',[
    'Unauthorised use or disclosure may cause harm not adequately compensated by damages. The Company may seek injunctive or equitable relief, in addition to other remedies, without waiving any right. Nothing in this Agreement excludes liability that cannot lawfully be excluded, limits protected employee rights, or predetermines damages contrary to applicable law.'
])

add_clause(doc,'17. General terms',[
    ('Notices.','Formal notices must be written and delivered to the addresses above (and to [LEGAL NOTICE EMAIL]) by personal delivery, reputable courier, or email with confirmation of receipt.'),
    ('Assignment.','The Team Member may not assign this Agreement without the Company’s prior written consent. The Company may assign it in connection with a reorganisation, financing, merger, acquisition or transfer of the EngiCite business.'),
    ('Entire agreement.','This Agreement and its schedules are the entire agreement on its subject and replace prior discussions on that subject, but do not reduce stricter duties in another signed agreement or customer requirement.'),
    ('Changes and waiver.','Changes must be in writing signed by authorised representatives. Delay or failure to enforce a right is not a waiver.'),
    ('Severability.','An invalid provision will be modified to the minimum extent needed to make it enforceable, and the remainder continues in effect.'),
    ('Counterparts and e-signatures.','This Agreement may be signed in counterparts and by electronic signature, each treated as an original.'),
    ('Order of precedence.','If a schedule conflicts with the body, the body controls unless the schedule expressly identifies the clause it overrides.'),
])

add_clause(doc,'18. Governing law and disputes',[
    'This Agreement is governed by the laws of [JURISDICTION], excluding conflict-of-law rules. The Parties submit to [EXCLUSIVE COURTS OF LOCATION / ARBITRATION RULES, SEAT, LANGUAGE AND NUMBER OF ARBITRATORS]. Either Party may seek urgent interim or injunctive relief from a court of competent jurisdiction. Complete this clause with qualified local counsel before use.'
])

doc.add_heading('Signatures',level=1)
doc.add_paragraph('Each signatory confirms that they have read, understood and agree to this Agreement and that they are authorised to sign for the identified Party.')
signature_block(doc,'For the Company')
signature_block(doc,'Team Member / Service Provider')

doc.add_page_break(); doc.add_heading('Schedule 1 — Engagement particulars',level=1)
fields=[('Engagement start date','[DATE]'),('Engagement end date (if fixed)','[DATE / NOT APPLICABLE]'),('Manager / sponsor','[NAME AND TITLE]'),('Approved role and team','[ROLE / TEAM]'),('Approved environments','[PRODUCTION / STAGING / DEVELOPMENT / NONE]'),('Approved customer/project scope','[LIST OR “AS ASSIGNED IN ENGICITE”]'),('Approved devices and repositories','[DETAILS]'),('Security incident contact','[EMAIL / PHONE]'),('Privacy contact','[EMAIL]'),('Legal notice contact','[EMAIL / ADDRESS]')]
t=doc.add_table(rows=0,cols=2); t.autofit=False
for a,b in fields:
    c=t.add_row().cells; c[0].width=Inches(2.35); c[1].width=Inches(4.15); shade(c[0],LIGHT); set_cell_text(c[0],a,True,NAVY); set_cell_text(c[1],b)

doc.add_heading('Confidentiality and security acknowledgements',level=2)
checks=['I completed required security, privacy, acceptable-use and product training.','I received only the access needed for my role and know how to request or surrender access.','I will use only approved devices, repositories, communication tools and AI services.','I understand tenant isolation and will not cross customer or project boundaries.','I know how and when to report a suspected incident.','I disclosed all Background Materials and potential conflicts in Schedule 2.']
for x in checks: doc.add_paragraph('☐ '+x)
p=doc.add_paragraph('Team Member initials: ____________________    Date: ____________________'); p.paragraph_format.space_before=Pt(10)

doc.add_heading('Schedule 2 — Background Materials and permitted third-party materials',level=1)
doc.add_paragraph('List all pre-existing tools, code, templates, inventions, datasets or other materials that may be used in the Engagement. If none, write “NONE”. Approval to use an item does not waive licence, security or provenance review.')
t=doc.add_table(rows=1,cols=4); t.autofit=False; hdr=t.rows[0].cells
for c,txt,w in zip(hdr,['Item / version','Owner','Licence / restrictions','Approved by / date'],[1.55,1.25,2.25,1.45]): c.width=Inches(w); shade(c,NAVY); set_cell_text(c,txt,True,'FFFFFF',9)
for _ in range(5):
    row=t.add_row().cells
    for c in row: set_cell_text(c,'\n')

doc.add_heading('Schedule 3 — Team or service-provider joinder',level=1)
doc.add_paragraph('Use this schedule where additional individuals join under an entity’s engagement. Each individual agrees personally to comply with Clauses 1–9, 12–14 and 17–18, and acknowledges that access may be withdrawn if this joinder is incomplete. Add pages as needed.')
t=doc.add_table(rows=1,cols=5); t.autofit=False; hdr=t.rows[0].cells
for c,txt,w in zip(hdr,['Full name','Role / project','Email','Signature','Date'],[1.45,1.45,1.55,1.3,.75]): c.width=Inches(w); shade(c,NAVY); set_cell_text(c,txt,True,'FFFFFF',8.5)
for _ in range(8):
    row=t.add_row().cells
    for c in row: set_cell_text(c,'\n')

doc.add_paragraph(); p=doc.add_paragraph('Document control'); font(p.add_run('Template version: 1.0  |  Prepared: 4 August 2026  |  Status: Counsel-review draft'),8,True,GRAY)

OUT.parent.mkdir(parents=True,exist_ok=True); doc.core_properties.title='EngiCite Non-Disclosure Agreement'; doc.core_properties.subject='Confidentiality, intellectual property and security agreement'; doc.core_properties.author='EngiCite'; doc.save(OUT)
print(OUT.resolve())
