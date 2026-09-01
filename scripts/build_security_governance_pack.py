from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "compliance" / "EngiCite_Security_Governance_Pack_v1.0.docx"
LOGO = ROOT / "apps" / "web" / "public" / "engicite-logo-transparent.png"

NAVY = "071B33"
NAVY_2 = "10243E"
ORANGE = "F97316"
GREEN = "0C5B45"
INK = "1D2A3A"
MUTED = "617083"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F1ED"
PALE_ORANGE = "FFF0E9"
LIGHT = "F5F7F9"
WHITE = "FFFFFF"
RED = "A5452F"
GOLD = "8A6200"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def set_font(run, name="Aptos", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D8E0E8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def paragraph_border(paragraph, color=ORANGE, size="14"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abs_base = max(existing_abs or [0]) + 1
    num_base = max(existing_num or [0]) + 1

    def abstract(abs_id, num_fmt, text, font=None):
        node = OxmlElement("w:abstractNum")
        node.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        node.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        node.append(lvl)
        first_num = numbering.find(qn("w:num"))
        if first_num is None:
            numbering.append(node)
        else:
            numbering.insert(list(numbering).index(first_num), node)

    def instance(num_id, abs_id):
        node = OxmlElement("w:num")
        node.set(qn("w:numId"), str(num_id))
        abs_num = OxmlElement("w:abstractNumId")
        abs_num.set(qn("w:val"), str(abs_id))
        node.append(abs_num)
        numbering.append(node)

    abstract(abs_base, "bullet", "•", "Arial")
    instance(num_base, abs_base)
    abstract(abs_base + 1, "decimal", "%1.")
    instance(num_base + 1, abs_base + 1)

    # Explicitly align the built-in List Bullet definition to the compact
    # reference-guide token map. Word's built-in style is retained because it
    # renders reliably across Word and LibreOffice.
    for node in numbering.findall(qn("w:abstractNum")):
        if node.get(qn("w:abstractNumId")) != "8":
            continue
        lvl = node.find(qn("w:lvl"))
        p_pr = lvl.find(qn("w:pPr"))
        tabs = p_pr.find(qn("w:tabs"))
        tab = tabs.find(qn("w:tab"))
        tab.set(qn("w:pos"), "540")
        ind = p_pr.find(qn("w:ind"))
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
    return 1, abs_base + 1


def new_number_sequence(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing or [0]) + 1
    node = OxmlElement("w:num")
    node.set(qn("w:numId"), str(num_id))
    abs_num = OxmlElement("w:abstractNumId")
    abs_num.set(qn("w:val"), str(abstract_id))
    node.append(abs_num)
    # Word can continue numbering across separate num instances that share an
    # abstract definition. Force every procedure to restart at step 1.
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    node.append(lvl_override)
    numbering.append(node)
    return num_id


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, NAVY_2, 14, 7),
        ("Heading 3", 11.5, GREEN, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, fill, color in (
        ("Policy Callout", PALE_BLUE, NAVY),
        ("Action Callout", PALE_GREEN, GREEN),
        ("Warning Callout", PALE_ORANGE, RED),
    ):
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.right_indent = Inches(0.18)
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(8)
        style._engicite_fill = fill

    small = doc.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    small.base_style = normal
    small.font.size = Pt(8.5)
    small.font.color.rgb = RGBColor.from_string(MUTED)
    small.paragraph_format.space_after = Pt(4)

    table_text = doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    table_text.base_style = normal
    table_text.font.size = Pt(8.6)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.05


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), ORANGE)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_callout(doc, text, style="Policy Callout"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    fill = {"Policy Callout": PALE_BLUE, "Action Callout": PALE_GREEN, "Warning Callout": PALE_ORANGE}[style]
    shade_paragraph(p, fill)
    return p


def add_bullet(doc, text, bullet_num_id, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=NAVY_2)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_numbered(doc, text, number_num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(number_num_id))
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text="", *, bold_prefix=None, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=NAVY_2)
        p.add_run(text[len(bold_prefix):])
    else:
        r = p.add_run(text)
        r.italic = italic
    return p


def add_table(doc, headers, rows, widths_dxa, header_fill=NAVY_2, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    header = table.rows[0]
    repeat_header(header)
    prevent_row_split(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.style = "Table Text"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_font(r, size=font_size, color=WHITE, bold=True)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(values):
            cell = cells[idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.style = "Table Text"
            p.add_run(str(value))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(2)
    after.paragraph_format.space_after = Pt(2)
    return table


def setup_headers_and_footers(doc):
    for idx, section in enumerate(doc.sections):
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        section.header_distance = Inches(0.36)
        section.footer_distance = Inches(0.35)
        section.different_first_page_header_footer = idx == 0
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("ENGICITE  |  SECURITY GOVERNANCE")
        set_font(r, size=8, color=MUTED, bold=True)
        paragraph_border(p, color="D7DFE7", size="6")
        footer = section.footer
        table = footer.add_table(rows=1, cols=2, width=Inches(6.86))
        repeat_header(table.rows[0])
        set_table_geometry(table, [7200, 2160], indent=0)
        set_table_borders(table, color=WHITE, size="0")
        p1 = table.cell(0, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run("INTERNAL - CONTROLLED  |  EGC-ISMS-GOV-001  |  v1.0")
        set_font(r1, size=8, color=MUTED)
        p2 = table.cell(0, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        add_page_number(p2)


def cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    picture = p.add_run().add_picture(str(LOGO), width=Inches(3.65))
    picture._inline.docPr.set("descr", "EngiCite corporate logo")
    picture._inline.docPr.set("title", "EngiCite")

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    r = kicker.add_run("SECURITY & TRUST PROGRAMME")
    set_font(r, size=10, color=ORANGE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run("Security Governance Pack")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    r = subtitle.add_run("Policies, responsibilities, registers and evidence requirements for SOC 2 readiness")
    set_font(r, size=14, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    paragraph_border(rule, ORANGE, "18")

    meta = [
        ("Document ID", "EGC-ISMS-GOV-001"),
        ("Version", "1.0"),
        ("Status", "DRAFT FOR APPROVAL"),
        ("Classification", "Internal - Controlled"),
        ("Policy owner", "Information Security Officer (Founder acting on an interim basis)"),
        ("Approver", "Founder / Board or delegated executive"),
        ("Proposed effective date", "Upon formal approval"),
        ("Review cycle", "At least annually and after material security change"),
    ]
    add_table(doc, ["Control", "Value"], meta, [2300, 7060], header_fill=NAVY)
    add_callout(doc, "Approval notice: this pack becomes an operating control only after the approval record in Appendix I is completed. Until then, it is a controlled implementation draft.", "Warning Callout")


def front_matter(doc, bullet_id):
    doc.add_page_break()
    add_heading(doc, "Document control", 1)
    add_table(doc, ["Version", "Date", "Author / owner", "Change", "Approval"], [
        ("1.0", "31 Aug 2026", "EngiCite Founder / Interim Security Officer", "Initial consolidated security governance pack", "Pending"),
    ], [900, 1200, 2250, 3710, 1300])

    add_heading(doc, "How to use this pack", 2)
    add_para(doc, "This document establishes the minimum security rules for EngiCite and provides the initial registers and evidence forms needed to operate those rules. Policy statements use the word must where compliance is mandatory. Any exception must follow the exception process in Section 18.")
    for item in (
        "Approve the pack and appoint named control owners.",
        "Complete the initial risk, asset and vendor registers; remove unresolved placeholders.",
        "Operate each recurring control at the stated frequency and retain evidence.",
        "Review the pack at least annually and after material changes, incidents or legal requirements.",
        "Keep customer documents and secrets out of the governance evidence repository.",
    ):
        add_bullet(doc, item, bullet_id)

    add_heading(doc, "Contents", 2)
    contents = [
        "1. Governance charter and scope", "2. Security principles and objectives", "3. Roles and accountability",
        "4. Information security policy", "5. Risk management", "6. Asset and data classification",
        "7. Identity and access control", "8. Secure development and change management",
        "9. Vulnerability and patch management", "10. Logging, monitoring and audit",
        "11. Incident and breach response", "12. Vendor and subprocessor security",
        "13. Availability, continuity and backup", "14. Data retention and secure disposal",
        "15. Personnel security and awareness", "16. Physical and endpoint security",
        "17. Customer security and communications", "18. Exceptions, enforcement and review",
        "19. Control operating calendar", "Appendices A-I: registers, forms and approval record",
    ]
    table_rows = [(contents[i], contents[i + 1] if i + 1 < len(contents) else "") for i in range(0, len(contents), 2)]
    add_table(doc, ["Governance sections", "Governance sections"], table_rows, [4680, 4680], header_fill=GREEN)


def governance_scope(doc, bullet_id):
    add_heading(doc, "1. Governance charter and scope", 1)
    add_callout(doc, "Security objective: protect the confidentiality, integrity and availability of EngiCite, customer engineering records, identities, audit evidence and business operations.")
    add_heading(doc, "1.1 Purpose", 2)
    add_para(doc, "This pack defines EngiCite's information security governance system. It converts technical protections and operational practices into assigned, reviewable and evidence-producing controls suitable for customer assurance and SOC 2 readiness.")
    add_heading(doc, "1.2 Organisational scope", 2)
    add_para(doc, "The policy applies to EngiCite's founders, directors, employees, contractors, temporary workers, service providers and any person granted access to EngiCite information or production systems.")
    add_heading(doc, "1.3 System boundary", 2)
    for item in (
        "EngiCite web application and APIs hosted on Vercel.",
        "Supabase authentication, PostgreSQL database, Row Level Security and object storage.",
        "FastAPI document processor, ClamAV scanning and supporting runtime hosted on Railway.",
        "GitHub source repositories, CI/CD workflows and deployment integrations.",
        "OpenAI APIs used for embeddings and evidence-grounded answers.",
        "Resend transactional email, Paystack billing, DNS/domain services and approved backup destinations.",
        "Administrator endpoints, work devices, security records, support tools and customer-facing operational processes.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "1.4 Out of scope", 2)
    add_para(doc, "Personal systems, unapproved applications and customer-controlled systems are outside EngiCite's direct control. They remain subject to contractual interfaces, vendor review and documented shared-responsibility boundaries.")


def principles(doc, bullet_id):
    add_heading(doc, "2. Security principles and objectives", 1)
    principles_data = [
        ("Tenant isolation", "Every customer organisation and project is isolated by server-side checks and database policies."),
        ("Least privilege", "Users and services receive only the minimum access needed for assigned duties."),
        ("Defence in depth", "Identity, application, database, storage, processor and monitoring controls reinforce one another."),
        ("Secure by default", "Private storage, no-store responses, restricted roles and fail-closed processing are the default."),
        ("Traceability", "Material actions must be attributable, timestamped and reviewable."),
        ("Data minimisation", "Only data required for the service and contractual purpose may be collected or disclosed."),
        ("Resilience", "Backups, recovery, monitoring and incident procedures must support approved RPO and RTO."),
        ("Continuous improvement", "Risks, incidents, tests and customer feedback must drive measurable control improvements."),
    ]
    add_table(doc, ["Principle", "EngiCite requirement"], principles_data, [2100, 7260], header_fill=NAVY)
    add_heading(doc, "2.1 Annual security objectives", 2)
    for item in (
        "No unauthorised cross-tenant access and no unapproved public customer-document storage.",
        "All privileged workforce and application roles protected by MFA.",
        "All production changes reviewed, tested, approved and traceable.",
        "All critical and high vulnerabilities remediated within the approved service levels.",
        "Quarterly access reviews, restore tests and security control reviews completed without overdue actions.",
        "At least one independent penetration test and one incident-response exercise each year.",
    ):
        add_bullet(doc, item, bullet_id)


def roles(doc):
    add_heading(doc, "3. Roles and accountability", 1)
    add_para(doc, "One person may temporarily hold multiple roles in an early-stage company, but the action and independent review responsibilities must remain separate for high-risk changes, access grants, deletions and risk acceptance.")
    rows = [
        ("Founder / Board", "Approves policy, risk appetite, major exceptions, resources and SOC 2 scope."),
        ("Information Security Officer", "Owns the programme, risk register, incidents, evidence, access reviews and annual reporting."),
        ("Engineering Lead", "Owns secure design, code review, CI/CD, dependency, vulnerability and remediation controls."),
        ("Operations Lead", "Owns monitoring, availability, backups, recovery, provider configuration and operational evidence."),
        ("Privacy / Legal Owner", "Owns privacy notice, retention, contracts, data requests, subprocessors and breach obligations."),
        ("System Owner", "Approves business access, classifies data, reviews risks and validates control operation."),
        ("Independent Reviewer", "Reviews high-risk changes, access certifications, deletions, exceptions and control evidence."),
        ("All Personnel", "Protect credentials and information, complete training, follow policy and report concerns immediately."),
    ]
    add_table(doc, ["Role", "Accountability"], rows, [2400, 6960], header_fill=GREEN)
    add_heading(doc, "3.1 Minimum segregation of duties", 2)
    add_table(doc, ["High-risk activity", "Performer", "Independent approver / reviewer"], [
        ("Production code change", "Author / engineer", "Different authorised reviewer"),
        ("Privileged access grant", "System administrator", "Security or system owner"),
        ("Tenant physical deletion", "Authorised operator", "Second authorised approver"),
        ("Critical-risk acceptance", "Risk owner", "Founder / Board"),
        ("Emergency production change", "On-call engineer", "Post-change reviewer within one business day"),
        ("Backup restoration", "Operations operator", "System owner validates completeness"),
    ], [2750, 2600, 4010], header_fill=NAVY_2)


def info_security_policy(doc, bullet_id):
    add_heading(doc, "4. Information security policy", 1)
    add_callout(doc, "EngiCite must protect information according to its sensitivity, contractual commitments and operational importance throughout its lifecycle.")
    sections = {
        "4.1 Governance requirements": [
            "Security risks must be reviewed at least quarterly and before material architectural or vendor change.",
            "Control owners must retain evidence that recurring controls operated at the required frequency.",
            "Security requirements must be included in product design, procurement, contracts and release decisions.",
            "Known control failures must be recorded, assessed, assigned and tracked to closure.",
        ],
        "4.2 Data protection requirements": [
            "Customer documents must remain in approved private storage and be accessed through authorised application paths.",
            "Customer content, credentials, tokens and secrets must not be placed in source code, tickets, screenshots, chat messages or test fixtures.",
            "Encryption in transit must be used for external and administrative connections; provider encryption at rest must be enabled.",
            "Downloads, exports and backups must use controlled access, short-lived links and appropriate encryption.",
        ],
        "4.3 Service protection requirements": [
            "Production and non-production environments must be separated by credentials and access boundaries.",
            "Internet-facing services must apply request validation, authentication, rate limits and security monitoring proportionate to risk.",
            "Document processing must fail closed when malware scanning or integrity verification cannot complete.",
            "Security controls must not be disabled in production without a time-limited, approved exception.",
        ],
    }
    for heading, items in sections.items():
        add_heading(doc, heading, 2)
        for item in items:
            add_bullet(doc, item, bullet_id)


def risk_policy(doc, bullet_id, number_id):
    add_heading(doc, "5. Risk management", 1)
    add_heading(doc, "5.1 Method", 2)
    add_para(doc, "Risks are scored using likelihood and impact from 1 to 5. Inherent risk is assessed before controls; residual risk is assessed after existing controls. The score is likelihood multiplied by impact.")
    add_table(doc, ["Score", "Rating", "Required response"], [
        ("1-4", "Low", "Accept and monitor through normal control review."),
        ("5-9", "Moderate", "Assign an owner and planned treatment; review quarterly."),
        ("10-16", "High", "Create a dated treatment plan; executive review is required."),
        ("17-25", "Critical", "Immediate containment; Founder approval required for any temporary continued operation."),
    ], [1250, 1700, 6410], header_fill=NAVY)
    add_heading(doc, "5.2 Risk lifecycle", 2)
    sequence_id = new_number_sequence(doc, number_id)
    for step in (
        "Identify the asset, threat, weakness, consequence and affected customer or business commitment.",
        "Score inherent likelihood and impact with a stated rationale.",
        "Identify preventive, detective and corrective controls and their evidence.",
        "Score residual risk and select mitigation, transfer, avoidance or formal acceptance.",
        "Assign an accountable owner, due date and measurable treatment action.",
        "Review critical risks monthly, high risks quarterly and all other risks at least annually.",
        "Close a risk only after evidence confirms treatment and the owner approves the residual position.",
    ):
        add_numbered(doc, step, sequence_id)
    add_heading(doc, "5.3 Risk acceptance", 2)
    add_para(doc, "Critical residual risk may be accepted only by the Founder / Board for a maximum of 30 days. High residual risk requires Founder approval and expires within 90 days. Moderate and low risks may be accepted by the Security Officer and relevant system owner. Every acceptance must identify compensating controls and an expiry date.")


def asset_data_policy(doc, bullet_id):
    add_heading(doc, "6. Asset and data classification", 1)
    add_heading(doc, "6.1 Asset management", 2)
    for item in (
        "Every production system, repository, provider account, domain, secret store, endpoint and critical dataset must have an owner.",
        "The asset register must record purpose, provider, environment, criticality, classification, access method and review date.",
        "Unapproved software or storage must not process EngiCite or customer information.",
        "Assets must be reviewed quarterly and on acquisition, material change, transfer or disposal.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "6.2 Classification scheme", 2)
    add_table(doc, ["Class", "Examples", "Handling rule"], [
        ("Public", "Approved website and marketing material", "May be shared after authorised publication review."),
        ("Internal", "Policies, internal plans and ordinary operating records", "Access limited to personnel with a business need."),
        ("Confidential", "Customer metadata, audit logs, contracts and internal security records", "Authenticated least-privilege access; controlled transfer and retention."),
        ("Restricted", "Customer document content, secrets, service-role keys, malware samples, encryption keys", "Strongest access controls; no unapproved copying; encrypted transfer and tightly controlled evidence."),
    ], [1500, 3250, 4610], header_fill=GREEN)
    add_callout(doc, "Default classification: all customer engineering documents and native files are Restricted unless a customer-approved rule states otherwise.", "Warning Callout")


def access_policy(doc, bullet_id, number_id):
    add_heading(doc, "7. Identity and access control", 1)
    add_heading(doc, "7.1 Core rules", 2)
    rules = (
        "Every user must have a unique identity and belong to an active organisation before workspace access is granted.",
        "Shared administrator accounts are prohibited except a documented emergency account with controlled credentials and activity review.",
        "MFA is mandatory for Founder, Organisation Administrator, Project Manager, Document Controller and all workforce provider-console accounts.",
        "Access must be role-based, tenant-scoped, least privilege and approved by the relevant owner.",
        "Service credentials must be non-human, stored in approved secret managers, scoped by environment and rotated on exposure or personnel change.",
        "Role preview or support access must be read-only, time-bound, clearly displayed and audited.",
    )
    for item in rules:
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "7.2 Joiner, mover and leaver procedure", 2)
    sequence_id = new_number_sequence(doc, number_id)
    for step in (
        "Verify the requester's authority and the user's organisation, project, role and discipline assignment.",
        "Grant the minimum role using an expiring invitation; never transmit passwords.",
        "Verify email ownership and MFA before privileged access is activated.",
        "Review access when role, project, discipline, employer or contract changes.",
        "Disable access immediately on termination, confirmed compromise or loss of legitimate business need.",
        "Revoke sessions, rotate affected shared secrets and preserve the access-removal evidence.",
    ):
        add_numbered(doc, step, sequence_id)
    add_heading(doc, "7.3 Access review", 2)
    add_para(doc, "The Security Officer and system owner must review privileged access quarterly and all workforce/provider access at least quarterly. Customer project managers remain responsible for project-team membership. Reviews must identify dormant accounts, excessive roles, unresolved invitations, preview activity and emergency access.")
    add_heading(doc, "7.4 Authentication baseline", 2)
    add_table(doc, ["Control", "Minimum requirement"], [
        ("Passwords", "At least 12 characters; provider leaked-password protection enabled where available."),
        ("MFA", "Phishing-resistant method preferred; required for privileged and provider-console access."),
        ("Sessions", "Secure cookies, inactivity and expiry controls appropriate to risk; revoke after security events."),
        ("Recovery", "Verified recovery channels; no support disclosure of credentials; recovery events logged."),
        ("Rate limits", "Authentication and recovery endpoints protected against automation and enumeration."),
    ], [2250, 7110], header_fill=NAVY_2)


def secure_development(doc, bullet_id, number_id):
    add_heading(doc, "8. Secure development and change management", 1)
    add_heading(doc, "8.1 Standard change workflow", 2)
    sequence_id = new_number_sequence(doc, number_id)
    for step in (
        "Record the business purpose, affected systems, security impact and rollback or forward-fix plan.",
        "Develop in a non-production environment with no production customer documents in test fixtures.",
        "Run type, lint, unit, integration, tenant-isolation, secret, dependency and build checks.",
        "Obtain approval from a reviewer who did not author the high-risk change.",
        "Merge only through a protected branch after required checks pass.",
        "Deploy using an authorised automated path and verify health, security controls and data integrity.",
        "Record the deployment, approver, test evidence, exceptions and post-deployment result.",
    ):
        add_numbered(doc, step, sequence_id)
    add_heading(doc, "8.2 Mandatory engineering controls", 2)
    for item in (
        "Protected production branch; force-push and branch deletion disabled.",
        "Independent approval and passing security checks required before merge.",
        "Additive, reviewed database migrations with forward-fix and rollback guidance.",
        "Secrets prohibited from repositories, logs, client bundles and examples.",
        "Dependency lockfiles, image scanning, SBOM generation and documented vulnerability response.",
        "Security-sensitive code includes negative tests for tenants, permissions, input abuse and failure paths.",
        "Production changes are traceable to a ticket, pull request, deployment and reviewer.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "8.3 Emergency changes", 2)
    add_para(doc, "An emergency change may bypass prior independent review only to contain an active incident or restore a critical service. The change must be minimal, logged, tested as far as practicable, approved by the incident commander, and independently reviewed within one business day.")


def vulnerability(doc, bullet_id):
    add_heading(doc, "9. Vulnerability and patch management", 1)
    add_para(doc, "EngiCite must identify vulnerabilities through dependency audits, secret scans, CodeQL/SAST, container and image scanning, provider advisories, authenticated testing, external penetration testing and responsible disclosure.")
    add_table(doc, ["Severity", "Target containment / remediation", "Escalation"], [
        ("Critical", "Contain immediately; remediate or formally mitigate within 72 hours", "Founder, Security and Engineering immediately"),
        ("High", "Remediate within 14 calendar days", "Weekly executive review until closed"),
        ("Moderate", "Remediate within 30 calendar days", "Security owner tracks to closure"),
        ("Low", "Remediate within 90 calendar days or documented release cycle", "Review during monthly vulnerability review"),
    ], [1400, 4780, 3180], header_fill=NAVY)
    for item in (
        "Internet-exposed critical findings must trigger immediate risk assessment and possible service restriction.",
        "Remediation must be retested; closure requires evidence, not only a code change.",
        "Expired remediation dates require a time-limited risk exception and compensating controls.",
        "At least one independent, authenticated, multi-tenant penetration test must occur annually and after major architecture change.",
    ):
        add_bullet(doc, item, bullet_id)


def logging_monitoring(doc, bullet_id):
    add_heading(doc, "10. Logging, monitoring and audit", 1)
    add_heading(doc, "10.1 Required events", 2)
    events = (
        "Successful and failed authentication, recovery, MFA and privileged access events.",
        "Organisation, project, membership, invitation, role and discipline changes.",
        "Uploads, downloads, secure previews, revision submissions, reviews and transmittals.",
        "AI questions, evidence selection and provider request identifiers without logging document content.",
        "Founder access, impersonation/preview, billing events, backup activity and destructive operations.",
        "Malware results, processing failures, service-role errors, rate-limit events and cross-tenant denials.",
        "Deployment, configuration, secret, firewall and provider administrator changes.",
    )
    for item in events:
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "10.2 Log protection", 2)
    add_para(doc, "Logs must be timestamped, access controlled, protected from unauthorised alteration and retained according to the approved retention schedule. Restricted content, passwords, tokens, full service errors and customer-document text must not be logged. Clock sources must be synchronised by the hosting platforms.")
    add_heading(doc, "10.3 Alert and review baseline", 2)
    add_table(doc, ["Alert / review", "Frequency", "Owner", "Evidence"], [
        ("Critical availability and security alerts", "Continuous", "Operations / Security on call", "Alert, acknowledgement and incident/ticket"),
        ("Failed malware scans and processing queues", "Daily", "Operations", "Dashboard or query review"),
        ("Privileged and founder activity", "Weekly", "Security Officer", "Reviewed audit extract"),
        ("Authentication abuse and cross-tenant denials", "Weekly", "Security Officer", "Trend and exceptions"),
        ("Audit-log completeness", "Monthly", "Security + Engineering", "Reconciliation record"),
        ("Provider administrator activity", "Monthly", "Security Officer", "Provider activity-log review"),
    ], [3220, 1350, 2100, 2690], header_fill=GREEN)


def incident_response(doc, bullet_id, number_id):
    add_heading(doc, "11. Incident and breach response", 1)
    add_heading(doc, "11.1 Reporting", 2)
    add_para(doc, "Personnel and suppliers must immediately report suspected unauthorised access, malware, data disclosure, lost credentials, tenant-isolation failure, service compromise, fraud, backup failure or material control failure to the designated security channel. Good-faith reporting is protected.")
    add_heading(doc, "11.2 Severity", 2)
    add_table(doc, ["Level", "Example", "Initial response"], [
        ("SEV-1", "Confirmed or likely customer-data exposure, active cross-tenant access, destructive compromise or widespread outage", "Immediate incident command; containment begins at once"),
        ("SEV-2", "Material degradation, suspected compromise, malware escape or important control failure", "Incident owner within one hour"),
        ("SEV-3", "Limited issue without confirmed customer impact", "Triage during the same business day"),
    ], [1250, 5480, 2630], header_fill=NAVY)
    add_heading(doc, "11.3 Response lifecycle", 2)
    sequence_id = new_number_sequence(doc, number_id)
    for step in (
        "Detect and record the time, reporter, affected systems and available evidence.",
        "Classify severity and appoint an incident commander and scribe.",
        "Contain access, sessions, credentials, deployments or processing while preserving evidence.",
        "Investigate the timeline, root cause, affected tenants, records and provider dependencies.",
        "Eradicate the cause, rotate secrets, patch vulnerabilities and validate tenant boundaries.",
        "Recover through approved restoration or deployment and monitor for recurrence.",
        "Notify customers, regulators, insurers and providers according to contracts and legal advice.",
        "Complete a blameless post-incident review, actions and control updates within five business days for SEV-1/2.",
    ):
        add_numbered(doc, step, sequence_id)
    add_heading(doc, "11.4 Evidence and communications", 2)
    add_para(doc, "Only the incident commander or authorised communications owner may issue external statements. Incident evidence must be stored in a restricted location with an access log. Customer notices must be accurate, approved and consistent with contractual and legal timelines.")


def vendor_security(doc, bullet_id, number_id):
    add_heading(doc, "12. Vendor and subprocessor security", 1)
    add_heading(doc, "12.1 Due diligence", 2)
    sequence_id = new_number_sequence(doc, number_id)
    for step in (
        "Identify the service, data, access, locations, dependencies and business criticality.",
        "Review security documentation, independent assurance, incident history and shared-responsibility terms.",
        "Review the DPA, confidentiality, breach notice, retention, deletion, audit and subprocessor terms.",
        "Assess access control, encryption, availability, recovery and exit/export capabilities.",
        "Record residual risks and obtain approval before production data is processed.",
        "Reassess critical vendors annually and after material incidents or contractual changes.",
    ):
        add_numbered(doc, step, sequence_id)
    add_heading(doc, "12.2 Minimum requirements", 2)
    for item in (
        "Critical providers must maintain appropriate independent security assurance or an approved alternative review.",
        "Vendor access must be least privilege, time-limited where possible, and removed when no longer required.",
        "Customer document use for model training is prohibited unless explicitly and lawfully approved by EngiCite and the customer.",
        "Provider data-retention settings must match EngiCite commitments and be evidenced.",
        "An exit plan must allow secure export, transition and confirmed deletion.",
    ):
        add_bullet(doc, item, bullet_id)


def continuity_backup(doc, bullet_id, number_id):
    add_heading(doc, "13. Availability, continuity and backup", 1)
    add_callout(doc, "Approved service objective: recovery point objective (RPO) of 24 hours and recovery time objective (RTO) of 4 hours, subject to contract and validated recovery testing.", "Action Callout")
    add_heading(doc, "13.1 Backup requirements", 2)
    for item in (
        "Database backup/PITR and object-storage backup must be treated as separate controls.",
        "At least one encrypted backup copy must be held outside the primary provider or failure domain.",
        "Backup access must be restricted, logged and protected by separately controlled credentials or keys.",
        "Backup jobs must be monitored; failures must alert an owner and be resolved within one business day.",
        "Backup retention and deletion must follow the approved retention matrix and legal holds.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "13.2 Recovery testing", 2)
    sequence_id = new_number_sequence(doc, number_id)
    for step in (
        "Select a representative tenant/project and record the approved test scope.",
        "Restore database records, storage objects and required application configuration into an isolated environment.",
        "Verify counts, checksums, revision links, permissions, audit history and critical workflows.",
        "Measure actual RPO and RTO and compare them with approved objectives.",
        "Document exceptions, corrective actions, owner and due date; obtain system-owner sign-off.",
    ):
        add_numbered(doc, step, sequence_id)
    add_para(doc, "A full restore test must be completed at least quarterly. A tabletop continuity exercise must occur at least annually and after major architecture change.")


def retention_disposal(doc, bullet_id):
    add_heading(doc, "14. Data retention and secure disposal", 1)
    add_para(doc, "EngiCite must retain information only for approved business, contractual, legal and security purposes. The Privacy / Legal Owner must approve the retention matrix before this policy becomes effective.")
    add_table(doc, ["Information type", "Proposed baseline", "Disposition control"], [
        ("Active customer documents and MDR records", "Contract term plus approved customer retention", "Tenant-authorised export and controlled deletion"),
        ("Application audit logs", "Minimum 12 months online; longer if contract or audit requires", "Restricted deletion; preserve legal holds"),
        ("Security and incident evidence", "Minimum 24 months or legal requirement", "Security-owner approval"),
        ("Backups", "Defined rolling period aligned with RPO and legal obligations", "Provider and off-provider deletion verification"),
        ("Invitations and failed onboarding", "Minimum period required for security and support", "Automated expiry and purge"),
        ("Billing records", "Applicable financial and tax requirement", "Restricted access and documented disposal"),
    ], [2880, 3000, 3480], header_fill=NAVY_2)
    add_heading(doc, "14.1 Disposal rules", 2)
    for item in (
        "Physical tenant deletion requires verified authority, two-person approval and a recorded cooling-off period unless lawfully overridden.",
        "Deletion must address database rows, Storage objects, search indexes, logs, backups and relevant vendors according to the retention matrix.",
        "Secrets and encryption keys must be revoked or destroyed when retired.",
        "Storage media and endpoints must use approved secure wipe or provider-certified destruction.",
        "Legal holds suspend ordinary deletion until formally released.",
    ):
        add_bullet(doc, item, bullet_id)


def personnel_endpoint(doc, bullet_id):
    add_heading(doc, "15. Personnel security and awareness", 1)
    for item in (
        "Security responsibilities must be included in contracts, onboarding and role descriptions.",
        "Personnel must sign confidentiality commitments before receiving Restricted access.",
        "Security awareness training is required on onboarding and annually, with role-specific secure-development training for engineers.",
        "Phishing, credential handling, customer-document handling, incident reporting and acceptable use must be covered.",
        "Policy violations are investigated consistently and may lead to access restriction, contract action or disciplinary measures.",
        "Offboarding must remove access, recover assets, revoke sessions and confirm ongoing confidentiality obligations.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "16. Physical and endpoint security", 1)
    for item in (
        "Work devices accessing production or Restricted data must use supported operating systems, disk encryption, screen lock, endpoint protection and automatic security updates.",
        "Local administrator privileges must be limited and reviewed.",
        "Restricted data must not be stored on unmanaged personal devices or removable media.",
        "Remote work must use trusted networks and approved secure connections; public devices are prohibited.",
        "Lost or stolen devices must be reported immediately and remotely disabled or wiped where supported.",
        "Physical work areas must prevent unauthorised viewing, photography or removal of customer records.",
    ):
        add_bullet(doc, item, bullet_id)


def customer_and_exceptions(doc, bullet_id, number_id):
    add_heading(doc, "17. Customer security and communications", 1)
    for item in (
        "Security claims must be accurate, approved and supported by evidence. EngiCite must not claim SOC 2 compliance before a CPA report is issued.",
        "Customer security questions must be answered from approved policies, current architecture and verified control evidence.",
        "Material security changes that affect contractual commitments must be communicated through the approved customer process.",
        "Customer data must not be used for model training; AI provider controls and retention must be disclosed accurately.",
        "Security incidents affecting customer information must follow approved contractual and legal notification procedures.",
    ):
        add_bullet(doc, item, bullet_id)
    add_heading(doc, "18. Exceptions, enforcement and review", 1)
    add_heading(doc, "18.1 Exception process", 2)
    sequence_id = new_number_sequence(doc, number_id)
    for step in (
        "State the control, reason, affected assets, data and duration.",
        "Assess likelihood, impact and residual risk.",
        "Define compensating controls, monitoring and an expiry date.",
        "Obtain the approval required by the risk rating.",
        "Record the exception in the risk register and review it before expiry.",
        "Close, renew through new approval, or remediate; exceptions must not renew automatically.",
    ):
        add_numbered(doc, step, sequence_id)
    add_heading(doc, "18.2 Enforcement", 2)
    add_para(doc, "Failure to follow this pack may result in removal of access, remediation requirements, contractual action or disciplinary measures. Deliberate concealment of incidents, unauthorised access or security-control bypass must be escalated immediately.")
    add_heading(doc, "18.3 Review", 2)
    add_para(doc, "The Security Officer must review this pack annually and after material incidents, architectural changes, acquisitions, regulatory changes or audit findings. The approver must authorise each new version before it becomes effective.")


def control_calendar(doc):
    add_heading(doc, "19. Control operating calendar", 1)
    rows = [
        ("Continuous", "Security and availability alerting; malware scanning; audit-event generation; signed access control", "Operations / Security"),
        ("Daily", "Backup and job failure review; processor health; critical advisories", "Operations"),
        ("Weekly", "Privileged/founder activity; auth abuse; cross-tenant denials; vulnerability exceptions", "Security"),
        ("Monthly", "Log review; vulnerability review; secret age; audit completeness; incident actions; backup success", "Security + Engineering + Operations"),
        ("Quarterly", "Access certification; risk register; asset/vendor register; restore test; RLS/security test evidence", "Control owners + independent reviewer"),
        ("Semi-annual", "Customer commitment review; continuity tabletop; data-flow and retention review", "Founder + Security + Legal"),
        ("Annual", "Policy approval; training; penetration test; vendor assurance; incident exercise; SOC 2 readiness review", "Founder / Board"),
        ("Event-driven", "Joiner/mover/leaver; incident; emergency change; major vendor/architecture change; tenant deletion", "Named process owner"),
    ]
    add_table(doc, ["Frequency", "Control activities", "Accountable owner"], rows, [1550, 5530, 2280], header_fill=ORANGE)
    add_callout(doc, "Evidence rule: if an activity is not recorded with date, owner, reviewer, result and exceptions, an auditor may treat it as not performed.", "Warning Callout")


def appendices(doc, bullet_id):
    doc.add_page_break()
    add_heading(doc, "Appendix A - Initial security risk register", 1)
    risks = [
        ("R-001", "Cross-tenant data access", "5x5 Critical", "RLS, server scope checks, tenant tests", "High", "Run SQL suite in CI/staging; quarterly RLS evidence", "Security / Engineering"),
        ("R-002", "Privileged account compromise", "4x5 Critical", "Founder AAL2, provider auth", "High", "Enforce MFA for all privileged roles and consoles", "Security"),
        ("R-003", "Unreviewed production change", "4x5 Critical", "CI, CODEOWNERS", "High", "Fix CI conflict; protected main; independent approval", "Engineering"),
        ("R-004", "Malicious document upload", "4x5 Critical", "Type/size checks, ClamAV fail closed", "Moderate", "Monitor scanner; annual upload-abuse test", "Operations"),
        ("R-005", "Backup or restoration failure", "4x5 Critical", "Portable checksummed backup", "High", "Off-provider encrypted copy; quarterly full restore", "Operations"),
        ("R-006", "Application or processor outage", "4x4 High", "Health checks, runbooks", "High", "Central alerts, capacity/load test, continuity exercise", "Operations"),
        ("R-007", "Secret exposure", "4x5 Critical", "Server-only keys, secret scan", "Moderate", "Rotation register; provider access review; detection", "Security / Engineering"),
        ("R-008", "Vendor retention or misuse", "4x5 Critical", "OpenAI store=false for answers", "High", "Vendor register, DPAs, ZDR/MAM evidence, annual review", "Legal / Security"),
        ("R-009", "Insufficient security logs", "4x4 High", "Application audit events", "High", "Central log drain, retention, alerts and monthly review", "Security / Operations"),
        ("R-010", "Bulk document exfiltration", "3x5 High", "RBAC, signed links, audit", "High", "Download anomaly alerts and response threshold", "Security"),
        ("R-011", "Invitation/account abuse", "3x4 High", "One-time tokens, role gates", "Moderate", "Rate limits, expiry review and alerting", "Engineering"),
        ("R-012", "Prompt injection or unsupported AI answer", "3x4 High", "Untrusted source prompt, citations, grounding", "Moderate", "Adversarial evals and monthly answer-quality sampling", "AI / Security"),
        ("R-013", "Billing webhook fraud or replay", "3x4 High", "HMAC, verification, deduplication", "Low", "Monthly billing-event reconciliation", "Finance / Engineering"),
        ("R-014", "Incomplete tenant deletion", "4x5 Critical", "Controlled purge functions", "High", "Approved retention matrix; deletion propagation test", "Legal / Operations"),
    ]
    add_table(doc, ["ID", "Risk", "Inherent", "Existing controls", "Residual", "Treatment", "Owner"], risks, [640, 1600, 900, 1800, 850, 2500, 1070], header_fill=NAVY, font_size=7.7)
    add_para(doc, "Register status: initial assessment. Owners must confirm scores, due dates and acceptance decisions before approval.", italic=True, after=8)

    add_heading(doc, "Appendix B - Initial asset register", 1)
    assets = [
        ("A-001", "EngiCite web/API", "Vercel", "Restricted", "Critical", "Engineering / Operations"),
        ("A-002", "PostgreSQL, Auth and Storage", "Supabase", "Restricted", "Critical", "Security / Operations"),
        ("A-003", "Document processor and malware scanner", "Railway", "Restricted", "Critical", "Engineering / Operations"),
        ("A-004", "Source and CI/CD", "GitHub", "Confidential", "Critical", "Engineering"),
        ("A-005", "AI answers and embeddings", "OpenAI", "Restricted", "High", "AI / Security"),
        ("A-006", "Transactional email", "Resend", "Confidential", "High", "Operations"),
        ("A-007", "Subscriptions and payments", "Paystack", "Confidential", "High", "Finance / Engineering"),
        ("A-008", "Domains and DNS", "Host / DNS provider", "Confidential", "Critical", "Operations"),
        ("A-009", "Portable project backups", "Supabase / approved external provider", "Restricted", "Critical", "Operations"),
        ("A-010", "Workforce endpoints", "EngiCite-managed", "Restricted", "High", "Security / individual custodian"),
    ]
    add_table(doc, ["ID", "Asset", "Provider / location", "Class", "Criticality", "Owner"], assets, [700, 2200, 2300, 1250, 1250, 1660], header_fill=GREEN)

    add_heading(doc, "Appendix C - Initial subprocessor and vendor register", 1)
    vendors = [
        ("Supabase", "Auth, database, storage", "Identity, metadata, documents", "Critical", "Obtain current SOC report/DPA; quarterly settings evidence"),
        ("Vercel", "Web hosting", "Requests, identity/session metadata", "Critical", "Obtain assurance; configure logs/firewall/access evidence"),
        ("Railway", "Document processing", "Temporary document content and processing metadata", "Critical", "Private service controls, logs, assurance and deletion"),
        ("OpenAI", "Embeddings and grounded answers", "Selected evidence text and questions", "Critical", "DPA, no-training statement, ZDR/MAM decision and evidence"),
        ("Resend", "Transactional email", "Email, organisation/project context", "High", "DPA, domain controls, retention and access review"),
        ("Paystack", "Billing", "Billing contact and transaction metadata", "High", "Contract, webhook controls and reconciliation"),
        ("GitHub", "Source and CI/CD", "Source, build logs, authorised identities", "Critical", "MFA, branch protection, assurance and access review"),
        ("DNS / domain provider", "Domain control", "Administrator identity and DNS records", "Critical", "MFA, registrar lock, recovery and access review"),
        ("FormSubmit", "Early-access form", "Prospect contact details", "Pending", "Remove or complete due diligence before production PII use"),
    ]
    add_table(doc, ["Provider", "Purpose", "Data", "Criticality", "Required evidence / decision"], vendors, [1350, 1800, 2380, 1100, 2730], header_fill=NAVY_2, font_size=8)

    add_heading(doc, "Appendix D - Quarterly access review record", 1)
    fields = [
        ("Review period", "____________________________"), ("Reviewer", "____________________________"),
        ("Independent approver", "____________________________"), ("Completion date", "____________________________"),
        ("Systems reviewed", "GitHub / Supabase / Vercel / Railway / Resend / Paystack / OpenAI / DNS / EngiCite roles"),
        ("Evidence location", "____________________________"),
    ]
    add_table(doc, ["Field", "Record"], fields, [2500, 6860], header_fill=GREEN)
    add_table(doc, ["Identity", "System / tenant", "Role", "MFA", "Business need confirmed", "Action / ticket"], [
        ("", "", "", "", "", ""), ("", "", "", "", "", ""), ("", "", "", "", "", ""),
    ], [1450, 1750, 1250, 700, 2200, 2010], header_fill=NAVY_2)

    add_heading(doc, "Appendix E - Security incident record", 1)
    incident_fields = [
        ("Incident ID / severity", ""), ("Detected / reported at", ""), ("Reporter and channel", ""),
        ("Incident commander / scribe", ""), ("Affected systems / tenants", ""), ("Data types and records", ""),
        ("Containment actions and times", ""), ("Evidence location and access", ""), ("Root cause", ""),
        ("Recovery and validation", ""), ("Notifications / legal decision", ""), ("Corrective actions / owners / dates", ""),
        ("Closure approval", ""),
    ]
    add_table(doc, ["Field", "Incident record"], incident_fields, [2800, 6560], header_fill=NAVY)

    add_heading(doc, "Appendix F - Production change approval record", 1)
    change_fields = [
        ("Change / ticket / pull request", ""), ("Purpose and affected systems", ""),
        ("Security and tenant impact", ""), ("Data migration / customer-file impact", ""),
        ("Test and security evidence", ""), ("Rollback / forward-fix plan", ""),
        ("Author", ""), ("Independent reviewer", ""), ("Deployment approver", ""),
        ("Deployment date / version", ""), ("Post-deployment validation", ""), ("Exceptions / follow-up", ""),
    ]
    add_table(doc, ["Field", "Change record"], change_fields, [2900, 6460], header_fill=GREEN)

    add_heading(doc, "Appendix G - Vendor security review checklist", 1)
    checks = [
        "Service, owner, criticality, data and locations identified",
        "Current independent assurance and bridge letter reviewed",
        "DPA, confidentiality, breach notice and deletion terms approved",
        "Subprocessors and international transfers reviewed",
        "MFA, access, encryption and audit controls assessed",
        "Availability, recovery, support and incident history assessed",
        "Retention, AI training use and exit/export capability assessed",
        "Residual risk, actions, approver and next review recorded",
    ]
    add_table(doc, ["Review criterion", "Result", "Evidence / action"], [(x, "Pass / Gap / N/A", "") for x in checks], [5200, 1500, 2660], header_fill=NAVY_2)

    add_heading(doc, "Appendix H - SOC 2 security evidence index", 1)
    evidence = [
        ("SEC-01", "Approved security policies and review history", "Annual", "Security Officer"),
        ("SEC-02", "Risk register review and treatment decisions", "Quarterly", "Security / risk owners"),
        ("SEC-03", "Privileged and provider access certification", "Quarterly", "Security / system owners"),
        ("SEC-04", "MFA and provider security configuration evidence", "Quarterly", "Security"),
        ("SEC-05", "Pull-request approval, CI and deployment sample", "Each release / sample", "Engineering"),
        ("SEC-06", "Vulnerability scans and remediation", "Monthly", "Engineering / Security"),
        ("SEC-07", "Penetration test and retest", "Annual", "Security"),
        ("SEC-08", "Alert tests and log review", "Monthly", "Security / Operations"),
        ("SEC-09", "Incident register and tabletop", "Event / annual", "Security"),
        ("SEC-10", "Backup and restoration test", "Quarterly", "Operations"),
        ("SEC-11", "Vendor assurance and DPA review", "Annual", "Legal / Security"),
        ("SEC-12", "Tenant isolation and database security tests", "Each release / quarterly", "Engineering"),
        ("SEC-13", "Training and policy acknowledgement", "Onboarding / annual", "People / Security"),
        ("SEC-14", "Exceptions and risk acceptance", "Event / monthly review", "Security / Founder"),
    ]
    add_table(doc, ["Evidence ID", "Evidence", "Frequency", "Owner"], evidence, [1150, 4900, 1500, 1810], header_fill=ORANGE)

    add_heading(doc, "Appendix I - Formal approval and appointments", 1)
    add_callout(doc, "Complete this page before treating the pack as an operating EngiCite policy.", "Warning Callout")
    appointments = [
        ("Operating legal entity", "_______________________________________________"),
        ("Founder / executive approver", "_______________________________________________"),
        ("Information Security Officer", "_______________________________________________"),
        ("Engineering control owner", "_______________________________________________"),
        ("Operations control owner", "_______________________________________________"),
        ("Privacy / Legal owner", "_______________________________________________"),
        ("Independent reviewer", "_______________________________________________"),
        ("Effective date", "_______________________________________________"),
        ("Next scheduled review", "_______________________________________________"),
    ]
    add_table(doc, ["Appointment / field", "Approved record"], appointments, [3500, 5860], header_fill=NAVY)
    add_para(doc, "Approval statement: I approve this Security Governance Pack as the minimum information-security policy for EngiCite. I accept responsibility for providing the resources, ownership and oversight necessary to operate the controls and retain reliable evidence.", after=14)
    add_table(doc, ["Approver name", "Signature", "Date"], [("", "", "")], [3500, 3500, 2360], header_fill=GREEN)

    add_heading(doc, "Authoritative references", 2)
    references = (
        "AICPA & CIMA, SOC 2 Trust Services Criteria: https://www.aicpa-cima.com/topic/audit-assurance/audit-and-assurance-greater-than-soc-2/",
        "Supabase, Shared Responsibility Model: https://supabase.com/docs/guides/deployment/shared-responsibility-model",
        "Supabase, Production Checklist: https://supabase.com/docs/guides/deployment/going-into-prod",
        "Vercel, Shared Responsibility Model: https://vercel.com/docs/security/shared-responsibility",
        "OpenAI, Data Controls in the OpenAI Platform: https://platform.openai.com/docs/models/default-usage-policies-by-endpoint",
        "EngiCite internal: SOC 2 Readiness Audit dated 31 August 2026; Threat Model; Incident Response; Backup/Restore; Monitoring and Release runbooks.",
    )
    for item in references:
        add_bullet(doc, item, bullet_id)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    configure_styles(doc)
    bullet_id, number_id = configure_numbering(doc)
    setup_headers_and_footers(doc)
    cover(doc)
    front_matter(doc, bullet_id)
    governance_scope(doc, bullet_id)
    principles(doc, bullet_id)
    roles(doc)
    info_security_policy(doc, bullet_id)
    risk_policy(doc, bullet_id, number_id)
    asset_data_policy(doc, bullet_id)
    access_policy(doc, bullet_id, number_id)
    secure_development(doc, bullet_id, number_id)
    vulnerability(doc, bullet_id)
    logging_monitoring(doc, bullet_id)
    incident_response(doc, bullet_id, number_id)
    vendor_security(doc, bullet_id, number_id)
    continuity_backup(doc, bullet_id, number_id)
    retention_disposal(doc, bullet_id)
    personnel_endpoint(doc, bullet_id)
    customer_and_exceptions(doc, bullet_id, number_id)
    control_calendar(doc)
    appendices(doc, bullet_id)

    core = doc.core_properties
    core.title = "EngiCite Security Governance Pack"
    core.subject = "SOC 2 security governance policies, responsibilities, registers and evidence requirements"
    core.author = "EngiCite"
    core.keywords = "EngiCite, security governance, SOC 2, access control, risk management"
    core.comments = "Controlled draft for approval"
    core.created = datetime(2026, 8, 31, 0, 0, 0)
    core.modified = datetime(2026, 8, 31, 0, 0, 0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
